#!/usr/bin/env python3
import argparse
import datetime as dt
import os
import platform
import subprocess
import re
import json
import shutil
import sys
import tempfile
import traceback
from pathlib import Path
from typing import Dict, Tuple

import xlrd
from xlutils.copy import copy as xl_copy

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk
except Exception:  # pragma: no cover
    tk = None
    filedialog = None
    messagebox = None
    ttk = None

IS_WINDOWS = platform.system().lower().startswith("win")


def detect_theme_mode() -> str:
    """Return 'dark' or 'light' based on Windows registry preference."""
    if IS_WINDOWS:
        try:
            import winreg  # type: ignore

            key = winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize",
            )
            apps_use_light, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
            return "light" if int(apps_use_light) == 1 else "dark"
        except Exception:
            pass
    return "light"


def build_theme_palette(mode: str) -> Dict[str, str]:
    if mode == "dark":
        return {
            "root_bg": "#0f172a",
            "header_bg": "#020617",
            "header_fg": "#e2e8f0",
            "header_sub_fg": "#94a3b8",
            "header_muted_fg": "#94a3b8",
            "header_active_bg": "#1e293b",
            "toolbar_fg": "#e2e8f0",
            "surface_bg": "#111827",
            "card_bg": "#1f2937",
            "card_border": "#334155",
            "section_fg": "#f59e0b",
            "label_fg": "#d1d5db",
            "entry_bg": "#0b1220",
            "entry_fg": "#e5e7eb",
            "entry_insert": "#f8fafc",
            "readonly_bg": "#1e293b",
            "readonly_fg": "#cbd5e1",
            "input_border": "#475569",
            "status_bg": "#020617",
            "status_fg": "#cbd5e1",
            "error_fg": "#fca5a5",
            "warn_fg": "#fde68a",
            "popup_select_bg": "#f59e0b",
            "popup_select_fg": "#0f172a",
            "btn_bg": "#1e293b",
            "btn_fg": "#cbd5e1",
            "scrollbar_bg": "#334155",
            "scrollbar_trough": "#1e293b",
            "scrollbar_arrow": "#94a3b8",
        }
    return {
        "root_bg": "#f4efe7",
        "header_bg": "#111827",
        "header_fg": "#ffffff",
        "header_sub_fg": "#cbd5e1",
        "header_muted_fg": "#9ca3af",
        "header_active_bg": "#1f2937",
        "toolbar_fg": "#ffffff",
        "surface_bg": "#f4efe7",
        "card_bg": "#fffdf8",
        "card_border": "#d6c7ad",
        "section_fg": "#7c2d12",
        "label_fg": "#374151",
        "entry_bg": "#ffffff",
        "entry_fg": "#374151",
        "entry_insert": "#111827",
        "readonly_bg": "#eef2ff",
        "readonly_fg": "#374151",
        "input_border": "#d1d5db",
        "status_bg": "#111827",
        "status_fg": "#d1d5db",
        "error_fg": "#fecaca",
        "warn_fg": "#fde68a",
        "popup_select_bg": "#e07b39",
        "popup_select_fg": "#ffffff",
        "btn_bg": "#4b5563",
        "btn_fg": "#ffffff",
        "scrollbar_bg": "#9ca3af",
        "scrollbar_trough": "#e5e7eb",
        "scrollbar_arrow": "#374151",
    }


def contract_sumtext_plain(text: str) -> str:
    """Return amount-in-words without trailing currency phrase."""
    cleaned = re.sub(r"\s+", " ", str(text or "").strip())
    cleaned = re.sub(
        r"\s*(грн\.?|грив(?:ня|ні|ень))\s*0{1,2}\s*коп(?:\.|ійок)?\s*$",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    return cleaned.strip(" ,.")


def _resolve_office_cli() -> str:
    for name in ("soffice", "libreoffice", "triooffice", "openoffice"):
        found = shutil.which(name)
        if found:
            return found
    for p in (
        r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    ):
        if Path(p).exists():
            return p
    return ""


def _office_convert(src_path: Path, out_dir: Path, to_ext: str) -> "Path | None":
    office_cli = _resolve_office_cli()
    if not office_cli:
        return None
    ext = to_ext.lstrip(".").lower()
    try:
        result = subprocess.run(
            [office_cli, "--headless", "--convert-to", ext, "--outdir", str(out_dir), str(src_path)],
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 0:
            return None
        out_path = out_dir / f"{src_path.stem}.{ext}"
        return out_path if out_path.exists() else None
    except Exception:
        return None


def _fill_docx_runs(para, values: Dict[str, str]) -> None:
    """Replace placeholders in paragraph runs, preserving run-level formatting."""
    for key, val in values.items():
        for ph in (f"{{{{{key}}}}}", f"<<{key}>>", f"[{key}]", f"${{{key}}}"):
            if ph not in para.text:
                continue
            # Strategy 1: single-run hit — preserves each run's own formatting
            for run in para.runs:
                if ph in run.text:
                    run.text = run.text.replace(ph, val)
            # Strategy 2: cross-run placeholder — merge all into first run (preserves its format)
            if ph in para.text:
                full = "".join(r.text for r in para.runs)
                if ph in full:
                    new_full = full.replace(ph, val)
                    if para.runs:
                        para.runs[0].text = new_full
                        for r in para.runs[1:]:
                            r.text = ""


def _remove_form_field_shading(doc) -> None:
    """Remove gray form-field backgrounds from a filled .docx (w:sdt controls + w:shd shading)."""
    try:
        from docx.oxml.ns import qn  # type: ignore
    except ImportError:
        return
    # Unwrap SDT content controls (structured document tags render as gray form boxes)
    for sdt in list(doc.element.body.iter(qn("w:sdt"))):
        sdt_content = sdt.find(qn("w:sdtContent"))
        if sdt_content is not None:
            parent = sdt.getparent()
            if parent is None:
                continue
            idx = list(parent).index(sdt)
            for child in list(sdt_content):
                parent.insert(idx, child)
                idx += 1
            parent.remove(sdt)
    # Remove run-level shading elements (w:shd causes gray highlight)
    for run_elem in list(doc.element.body.iter(qn("w:r"))):
        rpr = run_elem.find(qn("w:rPr"))
        if rpr is not None:
            for shd in list(rpr.findall(qn("w:shd"))):
                rpr.remove(shd)


def _fill_docx_contract_template(state: Dict[str, str], template_docx: Path, out_docx: Path) -> "Path | None":
    """Copy template byte-for-byte then fill ALL placeholders/bookmarks/content-controls."""
    try:
        from docx import Document  # type: ignore
        from docx.oxml.ns import qn as _qn  # type: ignore
    except Exception:
        return None

    if not template_docx.exists():
        return None

    payload = parse_state(state)
    values = {
        "Number": str(payload.get("Number", "")),
        "Data": str(payload.get("Data", "")),
        "FIO": str(payload.get("FIO", "")),
        "pasport": str(payload.get("pasport", "")),
        "TaxNumber": str(payload.get("TaxNumber", "")),
        "BirthDay": str(payload.get("BirthDay", "")),
        "adres": str(payload.get("adres", "")),
        "decl": str(payload.get("decl", "")),
        "model": str(payload.get("model", "")),
        "year": str(payload.get("year", "")),
        "color": str(payload.get("color", "")),
        "numberdv": str(payload.get("numberdv", "")),
        "cuzov": str(payload.get("cuzov", "")),
        "cub": str(payload.get("cub", "")),
        "znak": str(payload.get("znak", "")),
        "price": str(payload.get("price", "")),
        "sumtext": contract_sumtext_plain(payload.get("sumtext", "")),
        "FIO2": str(payload.get("FIO2", "")),
    }
    _XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    # Copy template byte-for-byte (preserves all styles, images, layout)
    shutil.copy2(template_docx, out_docx)
    doc = Document(str(out_docx))

    # --- Strategy A: fill SDT content controls by tag / alias name ----------
    # Modern Word templates often use structured document tags (gray form fields).
    # Match by w:tag or w:alias attribute so the correct field gets its value.
    for sdt in list(doc.element.body.iter(_qn("w:sdt"))):
        sdt_pr = sdt.find(_qn("w:sdtPr"))
        if sdt_pr is None:
            continue
        tag_el = sdt_pr.find(_qn("w:tag"))
        alias_el = sdt_pr.find(_qn("w:alias"))
        field_name = (
            (tag_el.get(_qn("w:val")) if tag_el is not None else None)
            or (alias_el.get(_qn("w:val")) if alias_el is not None else None)
        )
        if field_name and field_name in values:
            sdt_content = sdt.find(_qn("w:sdtContent"))
            if sdt_content is not None:
                all_t = list(sdt_content.iter(_qn("w:t")))
                if all_t:
                    all_t[0].text = values[field_name]
                    all_t[0].set(_XML_SPACE, "preserve")
                    for t in all_t[1:]:
                        t.text = ""

    # --- Strategy B: run-level fill on EVERY paragraph in the document -------
    # Covers both top-level paragraphs AND paragraphs nested inside content
    # controls (w:sdt) and table cells that doc.paragraphs / doc.tables miss.
    def _fill_para_elem(para_elem) -> None:
        runs = list(para_elem.iter(_qn("w:r")))
        if not runs:
            return
        for key, val in values.items():
            for ph in (f"{{{{{key}}}}}", f"<<{key}>>", f"[{key}]", f"${{{key}}}"):
                para_text = "".join((r.findtext(_qn("w:t")) or "") for r in runs)
                if ph not in para_text:
                    continue
                # Pass 1: single-run replacement (preserves run formatting)
                for run in runs:
                    t = run.find(_qn("w:t"))
                    if t is not None and t.text and ph in t.text:
                        t.text = t.text.replace(ph, val)
                        if t.text and (t.text[0] == " " or t.text[-1] == " "):
                            t.set(_XML_SPACE, "preserve")
                # Pass 2: cross-run merge into first run
                para_text_now = "".join((r.findtext(_qn("w:t")) or "") for r in runs)
                if ph in para_text_now:
                    new_text = para_text_now.replace(ph, val)
                    first_t = runs[0].find(_qn("w:t"))
                    if first_t is not None:
                        first_t.text = new_text
                        if new_text and (new_text[0] == " " or new_text[-1] == " "):
                            first_t.set(_XML_SPACE, "preserve")
                    for run in runs[1:]:
                        t = run.find(_qn("w:t"))
                        if t is not None:
                            t.text = ""

    for para_elem in doc.element.body.iter(_qn("w:p")):
        _fill_para_elem(para_elem)

    # --- Strategy C: fill Word bookmarks ------------------------------------
    for key, val in values.items():
        if not val:
            continue
        bookmarks = doc.element.xpath(f".//w:bookmarkStart[@w:name='{key}']")
        for bm in bookmarks:
            node = bm.getnext()
            while node is not None:
                text_nodes = node.xpath(".//w:t")
                if text_nodes:
                    text_nodes[0].text = val
                    for extra in text_nodes[1:]:
                        extra.text = ""
                    break
                node = node.getnext()

    _remove_form_field_shading(doc)
    doc.save(str(out_docx))
    return out_docx if out_docx.exists() else None


def generate_contract_via_office_fallback(state: Dict[str, str], template_doc: Path, out_path: Path) -> "Path | None":
    """Fallback without Word COM: try LibreOffice conversions while preserving template layout."""
    if not template_doc.exists():
        return None

    with tempfile.TemporaryDirectory(prefix="jm_contract_") as td:
        tmp_dir = Path(td)
        tmp_doc = tmp_dir / "template.doc"
        shutil.copy2(template_doc, tmp_doc)
        tpl_docx = _office_convert(tmp_doc, tmp_dir, "docx")
        if not tpl_docx:
            return None

        filled_docx = tmp_dir / "filled.docx"
        filled = _fill_docx_contract_template(state, tpl_docx, filled_docx)
        if not filled:
            return None

        if out_path.suffix.lower() == ".doc":
            converted_doc = _office_convert(filled, out_path.parent, "doc")
            if not converted_doc:
                return None
            final_doc = out_path.with_suffix(".doc")
            if converted_doc != final_doc:
                try:
                    if final_doc.exists():
                        final_doc.unlink()
                    converted_doc.rename(final_doc)
                except Exception:
                    return converted_doc
            return final_doc

        final_docx = out_path.with_suffix(".docx")
        shutil.copy2(filled, final_docx)
        return final_docx


def generate_contract_non_com(state: Dict[str, str], template_path: Path, out_path: Path) -> "Path | None":
    """Generate contract without Word COM.
    .docx template — fills directly (run-level, preserves formatting 1:1).
    .doc  template — converts via LibreOffice first.
    """
    suffix = template_path.suffix.lower()
    if suffix == ".docx":
        out_docx = out_path.with_suffix(".docx")
        filled = _fill_docx_contract_template(state, template_path, out_docx)
        if not filled:
            return None
        # Optionally convert to .doc if output extension requested
        if out_path.suffix.lower() == ".doc":
            doc_result = _office_convert(filled, out_path.parent, "doc")
            if doc_result:
                try:
                    filled.unlink()
                except Exception:
                    pass
                return doc_result
        return filled  # keep .docx — looks identical to original template
    # .doc template — use LibreOffice conversion workflow
    return generate_contract_via_office_fallback(state, template_path, out_path)


def runtime_app_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def runtime_resource_dir() -> Path:
    if getattr(sys, "frozen", False):
        meipass = getattr(sys, "_MEIPASS", None)
        if meipass:
            return Path(meipass).resolve()
    return Path(__file__).resolve().parent


def default_output_dir(base_dir: Path) -> Path:
    if IS_WINDOWS:
        return Path.home() / "Documents" / "JapanMoto" / "out"
    return base_dir / "out"


def safe_contract_output_path(out_path: Path) -> Path:
    return out_path.with_suffix(".txt")


def a1_to_rc(addr: str) -> Tuple[int, int]:
    m = re.fullmatch(r"([A-Z]+)(\d+)", addr.upper())
    if not m:
        raise ValueError(f"Invalid address: {addr}")
    col_s, row_s = m.groups()
    col = 0
    for ch in col_s:
        col = col * 26 + (ord(ch) - 64)
    return int(row_s) - 1, col - 1


def rc_to_a1(row: int, col: int) -> str:
    c = col + 1
    out = ""
    while c:
        c, rem = divmod(c - 1, 26)
        out = chr(65 + rem) + out
    return f"{out}{row + 1}"


def read_xls_cell(path: Path, sheet_name: str, addr: str):
    wb = xlrd.open_workbook(str(path), formatting_info=False)
    sh = wb.sheet_by_name(sheet_name)
    r, c = a1_to_rc(addr)
    return sh.cell_value(r, c)


def write_xls_cells(path: Path, sheet_name: str, updates: Dict[str, object], backup: bool = True) -> None:
    if backup:
        ts = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = path.with_name(f"{path.stem}_backup_{ts}{path.suffix}")
        shutil.copy2(path, backup_path)

    rb = xlrd.open_workbook(str(path), formatting_info=True)
    wb = xl_copy(rb)

    sheet_idx = None
    for i, name in enumerate(rb.sheet_names()):
        if name == sheet_name:
            sheet_idx = i
            break
    if sheet_idx is None:
        raise ValueError(f"Sheet not found: {sheet_name}")

    ws = wb.get_sheet(sheet_idx)
    rs = rb.sheet_by_index(sheet_idx)
    for addr, value in updates.items():
        r, c = a1_to_rc(addr)
        orig_xf = rs.cell_xf_index(r, c)
        ws.write(r, c, value)
        # Restore original XF index so cell formatting (fonts, borders, etc.) is preserved
        row_obj = ws._Worksheet__rows.get(r)
        if row_obj is not None:
            cell_obj = row_obj._Row__cells.get(c)
            if cell_obj is not None:
                cell_obj.xf_idx = orig_xf

    wb.save(str(path))


def split_number_date(text: str) -> Tuple[str, str]:
    # Example: "№ 8424/26/000308 від 15 травня 2026 року"
    parts = str(text).split()
    if len(parts) >= 7 and parts[0] == "№" and "від" in parts:
        number = parts[1]
        try:
            vidx = parts.index("від")
            date_txt = " ".join(parts[vidx + 1 : vidx + 5])
        except ValueError:
            date_txt = ""
        return number, date_txt

    # Fallback parser
    m = re.search(r"№\s*([^\s]+).*?від\s+(.+)$", str(text))
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return "", ""


def short_name(full_name: str) -> str:
    tokens = [t for t in str(full_name).strip().split() if t]
    if len(tokens) < 3:
        return str(full_name).strip()
    return f"{tokens[0]} {tokens[1][0]}. {tokens[2][0]}."


def load_source_values(source_6055: Path) -> Dict[str, object]:
    wb = xlrd.open_workbook(str(source_6055), formatting_info=False)
    sh = wb.sheet_by_name("Worksheet")

    def g(addr: str):
        r, c = a1_to_rc(addr)
        return sh.cell_value(r, c)

    a3 = str(g("A3"))
    number, date_txt = split_number_date(a3)
    fio = str(g("C15"))

    return {
        "number": number,
        "date": date_txt,
        "fio": fio,
        "fio_short": short_name(fio),
        "pasport": str(g("C17")),
        "tax": str(g("C18")),
        "birthday": str(g("C12")),
        "adres": str(g("C16")),
        "decl": str(g("C47")),
        "model": str(g("C20")),
        "year": str(g("C29")),
        "color": str(g("C28")),
        "numberdv": str(g("C41")),
        "cuzov": str(g("C39")),
        "cub": str(g("C36")),
        "znak": str(g("C50")),
        "price": str(g("C46")),
    }


def transfer_to_act_non_windows(source_6055: Path, moto_template: Path, out_path: Path) -> None:
    src = load_source_values(source_6055)

    shutil.copy2(moto_template, out_path)
    updates = {
        "B12": src["number"],
        "L12": src["date"],
        "H18": src["fio"],
        "C24": src["model"],
        "E24": src["year"],
        "H24": src["cuzov"],
        "K24": src["color"],
        "N33": src["fio_short"],
    }
    write_xls_cells(out_path, "Worksheet", updates, backup=False)


def transfer_to_word_preview(source_6055: Path, out_txt: Path) -> None:
    src = load_source_values(source_6055)
    lines = [
        "This is a Linux preview of Word bookmark payload.",
        "Use Windows mode to write into DOGOVIR_6055_template.doc bookmarks.",
        "",
    ]
    mapping = [
        ("Number", src["number"]),
        ("Data", src["date"]),
        ("FIO", src["fio"]),
        ("pasport", src["pasport"]),
        ("TaxNumber", src["tax"]),
        ("BirthDay", src["birthday"]),
        ("adres", src["adres"]),
        ("decl", src["decl"]),
        ("model", src["model"]),
        ("year", src["year"]),
        ("color", src["color"]),
        ("numberdv", src["numberdv"]),
        ("cuzov", src["cuzov"]),
        ("cub", src["cub"]),
        ("znak", src["znak"]),
        ("price", src["price"]),
        ("sumtext", "(computed by VBA in original workbook)"),
        ("FIO2", src["fio"]),
    ]
    for key, value in mapping:
        lines.append(f"{key}: {value}")

    out_txt.write_text("\n".join(lines), encoding="utf-8")


def windows_transfer_all(source_6055: Path, moto_template: Path, dogovir_template: Path, out_dir: Path) -> None:
    import win32com.client as win32  # type: ignore

    out_dir.mkdir(parents=True, exist_ok=True)
    excel = win32.Dispatch("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False

    word = win32.Dispatch("Word.Application")
    word.Visible = False

    try:
        wb_src = excel.Workbooks.Open(str(source_6055.resolve()))
        wb_dst = excel.Workbooks.Open(str(moto_template.resolve()))

        ws_src = wb_src.Worksheets("Worksheet")
        ws_dst = wb_dst.Worksheets("Worksheet")

        a3 = str(ws_src.Range("A3").Value)
        number, date_txt = split_number_date(a3)
        fio = str(ws_src.Range("C15").Value)

        ws_dst.Range("B12").Value = number
        ws_dst.Range("L12").Value = date_txt
        ws_dst.Range("H18").Value = fio
        ws_dst.Range("C24").Value = ws_src.Range("C20").Value
        ws_dst.Range("E24").Value = ws_src.Range("C29").Value
        ws_dst.Range("H24").Value = ws_src.Range("C39").Value
        ws_dst.Range("K24").Value = ws_src.Range("C28").Value
        ws_dst.Range("N33").Value = short_name(fio)

        ts = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        out_act = out_dir / f"6055_MOTO_filled_{ts}.xls"
        wb_dst.SaveAs(str(out_act.resolve()))

        doc = word.Documents.Open(str(dogovir_template.resolve()))
        bm = doc.Bookmarks

        payload = load_source_values(source_6055)
        payload["Number"] = payload.pop("number")
        payload["Data"] = payload.pop("date")
        payload["TaxNumber"] = payload.pop("tax")
        payload["BirthDay"] = payload.pop("birthday")
        payload["FIO"] = payload.pop("fio")
        payload["FIO2"] = payload["FIO"]

        bookmark_map = {
            "Number": "Number",
            "Data": "Data",
            "FIO": "FIO",
            "pasport": "pasport",
            "TaxNumber": "TaxNumber",
            "BirthDay": "BirthDay",
            "adres": "adres",
            "decl": "decl",
            "model": "model",
            "year": "year",
            "color": "color",
            "numberdv": "numberdv",
            "cuzov": "cuzov",
            "cub": "cub",
            "znak": "znak",
            "price": "price",
            "FIO2": "FIO2",
        }

        for key, bookmark_name in bookmark_map.items():
            if bm.Exists(bookmark_name):
                rng = bm(bookmark_name).Range
                rng.Text = str(payload[key])

        out_doc = out_dir / f"DOGOVIR_6055_filled_{ts}.doc"
        doc.SaveAs(str(out_doc.resolve()))
        doc.Close(SaveChanges=False)

        wb_dst.Close(SaveChanges=False)
        wb_src.Close(SaveChanges=True)
    finally:
        try:
            excel.Quit()
        except Exception:
            pass
        try:
            word.Quit()
        except Exception:
            pass


FIELD_SECTIONS = [
    (
        "Суб'єкт господарювання",
        [
            ("C5", "Найменування", False),
            ("C6", "Код ЄДРПОУ", False),
            ("C7", "Місцезнаходження", False),
            ("C8", "Реєстр. номер МВС", False),
            ("A56", "Продавець (підпис)", False),
        ],
    ),
    (
        "Документ",
        [
            ("A3", "Номер і дата", True),
            ("C47", "Митна декларація", False),
            ("C51", "Дата набуття права", False),
            ("C48", "Акт приймання-передачі", False),
            ("C49", "Свідоцтво про реєстрацію", False),
            ("C53", "Висновок", False),
        ],
    ),
    (
        "Транспорт",
        [
            ("C20", "Марка, модель", True),
            ("C21", "Тип ТЗ", False),
            ("C22", "Призначення", False),
            ("C23", "Сертифікат відповідності", False),
            ("C24", "Номер сертифіката типу", False),
            ("C25", "Категорія", False),
            ("C26", "Тип кузова", False),
            ("C27", "Паливо", False),
            ("C28", "Колір", True),
            ("C29", "Рік випуску", True),
            ("C30", "Повна маса", False),
            ("C31", "Маса без навантаження", False),
            ("C32", "Кількість дверей", False),
            ("C33", "Кількість місць", False),
            ("C34", "Кількість стоячих місць", False),
        ],
    ),
    (
        "",  # Ідентифікатори + Вартість — no section header
        [
            ("C35", "Кількість циліндрів", False),
            ("C36", "Об'єм двигуна", True),
            ("C37", "Потужність кВт", False),
            ("C38", "Потужність к.с.", False),
            ("C39", "VIN (номер рами)", True),
            ("C40", "Номер кузова", False),
            ("C41", "Номер двигуна", True),
            ("C42", "Номер шасі", True),
            ("C50", "Номерні знаки", False),
            ("C44", "Ціна без ПДВ", True),
            ("C45", "ПДВ", True),
            ("C46", "Ціна з ПДВ", True),
            ("DISC", "Знижка", False),
        ],
    ),
    (
        "Покупець",
        [
            ("C15", "ПІБ покупця", True),
            ("E15", "ПІБ скорочено", False),
            ("C12", "Дата народження", False),
            ("C16", "Адреса", True),
            ("C17", "Паспорт", True),
            ("C18", "ІПН / код", True),
        ],
    ),
]


def iter_field_specs():
    for group_name, fields in FIELD_SECTIONS:
        for cell, label, required in fields:
            yield group_name, cell, label, required


FORM_CELLS = [cell for _, cell, _, _ in iter_field_specs()]
CELL_LABELS: Dict[str, str] = {cell: label for _, cell, label, _ in iter_field_specs()}


def safe_read_cell(sheet, addr: str):
    try:
        r, c = a1_to_rc(addr)
        if r < 0 or c < 0 or r >= sheet.nrows or c >= sheet.ncols:
            return ""
        value = sheet.cell_value(r, c)
        if isinstance(value, float) and value.is_integer():
            return int(value)
        return value
    except Exception:
        return ""


def load_form_state(source_path: Path) -> Dict[str, str]:
    wb = xlrd.open_workbook(str(source_path), formatting_info=False)
    sh = wb.sheet_by_name("Worksheet")
    state: Dict[str, str] = {}
    for cell in FORM_CELLS:
        state[cell] = str(safe_read_cell(sh, cell)).strip()
    if not state.get("E15"):
        state["E15"] = short_name(state.get("C15", ""))
    return state


def parse_state(state: Dict[str, str]) -> Dict[str, str]:
    payload = dict(state)
    number, date_txt = split_number_date(payload.get("A3", ""))
    payload["Number"] = number
    payload["Data"] = date_txt
    payload["FIO"] = payload.get("C15", "")
    payload["FIO2"] = payload.get("C15", "")
    payload["TaxNumber"] = payload.get("C18", "")
    payload["BirthDay"] = payload.get("C12", "")
    payload["adres"] = payload.get("C16", "")
    payload["pasport"] = payload.get("C17", "")
    payload["decl"] = payload.get("C47", "")
    payload["model"] = payload.get("C20", "")
    payload["year"] = payload.get("C29", "")
    payload["color"] = payload.get("C28", "")
    payload["numberdv"] = payload.get("C41", "")
    payload["cuzov"] = payload.get("C39", "") or payload.get("C42", "") or payload.get("C43", "")
    payload["cub"] = payload.get("C36", "")
    payload["znak"] = payload.get("C50", "")
    payload["price"] = payload.get("C46", "")
    payload["sumtext"] = amount_to_words_uah(payload.get("C46", ""))
    payload["fio_short"] = short_name(payload.get("C15", ""))
    return payload


def amount_to_words_uah(value) -> str:
    try:
        amount = int(float(str(value).replace(",", ".")))
    except Exception:
        return ""

    if amount <= 0:
        return "нуль грн. 00 коп."

    ones = ["", "один", "два", "три", "чотири", "п'ять", "шість", "сім", "вісім", "дев'ять"]
    ones_f = ["", "одна", "дві", "три", "чотири", "п'ять", "шість", "сім", "вісім", "дев'ять"]
    teens = ["десять", "одинадцять", "дванадцять", "тринадцять", "чотирнадцять", "п'ятнадцять", "шістнадцять", "сімнадцять", "вісімнадцять", "дев'ятнадцять"]
    tens = ["", "", "двадцять", "тридцять", "сорок", "п'ятдесят", "шістдесят", "сімдесят", "вісімдесят", "дев'яносто"]
    hundreds = ["", "сто", "двісті", "триста", "чотириста", "п'ятсот", "шістсот", "сімсот", "вісімсот", "дев'ятсот"]

    def triad_words(num: int, feminine: bool = False) -> str:
        if num == 0:
            return ""
        parts = []
        h = num // 100
        t = (num // 10) % 10
        u = num % 10
        if h:
            parts.append(hundreds[h])
        if t == 1:
            parts.append(teens[u])
        else:
            if t:
                parts.append(tens[t])
            if u:
                parts.append((ones_f if feminine else ones)[u])
        return " ".join([p for p in parts if p])

    groups = [
        (amount % 1000, False, ""),
        ((amount // 1000) % 1000, True, "тисяч"),
        ((amount // 1000000) % 1000, False, "мільйонів"),
    ]

    chunks = []
    for idx, (value_part, feminine, suffix) in enumerate(groups):
        if not value_part:
            continue
        text = triad_words(value_part, feminine=feminine)
        if idx == 1:
            last_two = value_part % 100
            last = value_part % 10
            if 10 < last_two < 20:
                suffix = "тисяч"
            elif last == 1:
                suffix = "тисяча"
            elif last in (2, 3, 4):
                suffix = "тисячі"
            else:
                suffix = "тисяч"
        elif idx == 2:
            last_two = value_part % 100
            last = value_part % 10
            if 10 < last_two < 20:
                suffix = "мільйонів"
            elif last == 1:
                suffix = "мільйон"
            elif last in (2, 3, 4):
                suffix = "мільйона"
            else:
                suffix = "мільйонів"
        if suffix:
            chunks.append(f"{text} {suffix}".strip())
        else:
            chunks.append(text)

    return " ".join(reversed([chunk for chunk in chunks if chunk])).strip() + " грн. 00 коп."


def parse_decimal(value: str) -> float | None:
    try:
        cleaned = str(value).strip().replace(" ", "").replace(",", ".")
        if not cleaned:
            return None
        return float(cleaned)
    except Exception:
        return None


def normalize_token(value: str) -> str:
    return re.sub(r"[^A-Z0-9]", "", str(value).upper())


def valid_date_text(value: str) -> bool:
    value = str(value).strip()
    if not value:
        return False
    if re.fullmatch(r"\d{2}\.\d{2}\.\d{4}", value):
        return True
    return bool(re.search(r"\b\d{1,2}\b.+\b\d{4}\b", value))


def preview_blocks_for_act(payload: Dict[str, str]) -> list[tuple[str, list[tuple[str, str]]]]:
    return [
        ("Реквізити акта", [("Номер", payload.get("Number", "")), ("Дата", payload.get("Data", ""))]),
        (
            "Покупець",
            [
                ("ПІБ", payload.get("C15", "")),
                ("Адреса", payload.get("C16", "")),
                ("Паспорт", payload.get("C17", "")),
                ("ІПН", payload.get("C18", "")),
            ],
        ),
        (
            "Транспортний засіб",
            [
                ("Марка / модель", payload.get("C20", "")),
                ("Тип", payload.get("C21", "")),
                ("Рік", payload.get("C29", "")),
                ("Колір", payload.get("C28", "")),
                ("VIN / рама", payload.get("cuzov", "")),
                ("Номер двигуна", payload.get("C41", "")),
            ],
        ),
        ("Підписні дані", [("Скорочений ПІБ", payload.get("fio_short", "")), ("Дата набуття права", payload.get("C51", ""))]),
    ]


def preview_blocks_for_contract(payload: Dict[str, str]) -> list[tuple[str, list[tuple[str, str]]]]:
    return [
        ("Договір", [("Номер", payload.get("Number", "")), ("Дата", payload.get("Data", "")), ("Ціна", payload.get("price", "")), ("Сума прописом", payload.get("sumtext", ""))]),
        (
            "Покупець",
            [
                ("ПІБ", payload.get("FIO", "")),
                ("Дата народження", payload.get("BirthDay", "")),
                ("Паспорт", payload.get("pasport", "")),
                ("ІПН", payload.get("TaxNumber", "")),
                ("Адреса", payload.get("adres", "")),
            ],
        ),
        (
            "Мотоцикл",
            [
                ("Модель", payload.get("model", "")),
                ("Рік", payload.get("year", "")),
                ("Колір", payload.get("color", "")),
                ("Рама / кузов", payload.get("cuzov", "")),
                ("Двигун", payload.get("numberdv", "")),
                ("Об'єм", payload.get("cub", "")),
                ("Номерні знаки", payload.get("znak", "")),
            ],
        ),
        ("Митні дані", [("Декларація", payload.get("decl", ""))]),
    ]


def preview_blocks_for_vidatkova(payload: Dict[str, str]) -> list[tuple[str, list[tuple[str, str]]]]:
    return [
        ("Видаткова накладна", [("Дата / номер", payload.get("A3", "")), ("Одержувач", payload.get("C15", ""))]),
        (
            "Позиція",
            [
                ("Тип ТЗ", payload.get("C21", "")),
                ("Марка, модель", payload.get("C20", "")),
                ("Номер рами", payload.get("cuzov", "")),
                ("Кількість", "1"),
            ],
        ),
        (
            "Суми",
            [
                ("Ціна без ПДВ", payload.get("C44", "")),
                ("ПДВ", payload.get("C45", "")),
                ("Всього", payload.get("C46", "")),
                ("Сума прописом", payload.get("sumtext", "")),
            ],
        ),
    ]


def preview_blocks(kind: str, payload: Dict[str, str]) -> list[tuple[str, list[tuple[str, str]]]]:
    if kind == "act":
        return preview_blocks_for_act(payload)
    if kind == "contract":
        return preview_blocks_for_contract(payload)
    return preview_blocks_for_vidatkova(payload)


def validate_state(state: Dict[str, str]) -> Tuple[Dict[str, str], list[str], list[str]]:
    payload = parse_state(state)
    errors: list[str] = []
    warnings: list[str] = []
    invalid_cells: set[str] = set()
    warning_cells: set[str] = set()

    required = [cell for _, cell, _, req in iter_field_specs() if req]
    for cell in required:
        if not str(state.get(cell, "")).strip():
            label = CELL_LABELS.get(cell, cell)
            errors.append(f"«{label}» порожнє")
            invalid_cells.add(cell)

    if not payload["Number"] or not payload["Data"]:
        errors.append("Номер і дата документа (A3) не розпізнані")
        invalid_cells.add("A3")
    elif not re.fullmatch(r"\d{4}/\d{2}/\d{6}", payload["Number"]):
        warnings.append("Номер документа в A3 має незвичний формат")
        warning_cells.add("A3")

    if state.get("C51") and not valid_date_text(state.get("C51", "")):
        warnings.append("Дата набуття права — підозрілий формат")
        warning_cells.add("C51")

    if state.get("C12") and not valid_date_text(state.get("C12", "")):
        warnings.append("Дата народження — підозрілий формат")
        warning_cells.add("C12")

    frame_values = [str(state.get(c, "")).strip() for c in ("C39", "C42") if str(state.get(c, "")).strip()]
    if len(set(frame_values)) > 1:
        warnings.append("VIN / шасі не збігаються")
        warning_cells.update({"C39", "C42"})

    if payload.get("cuzov"):
        vin_token = normalize_token(payload["cuzov"])
        if len(vin_token) < 8:
            errors.append("VIN / рама занадто короткий")
            invalid_cells.add("C39")
        elif len(vin_token) not in (8, 9, 10, 11, 12, 13, 14, 15, 16, 17):
            warnings.append("VIN / рама має незвичну довжину")
            warning_cells.add("C39")

    if payload.get("C41"):
        engine_token = normalize_token(payload["C41"])
        if len(engine_token) < 5:
            warnings.append("Номер двигуна схожий на короткий або неповний")
            warning_cells.add("C41")
        if payload.get("cuzov") and engine_token == normalize_token(payload["cuzov"]):
            warnings.append("Номер двигуна збігається з VIN / рамою, перевірте дані")
            warning_cells.add("C41")

    if payload.get("C28") and payload["C28"].strip().isdigit():
        warnings.append("Колір виглядає як число")
        warning_cells.add("C28")

    if payload.get("C18"):
        tax = re.sub(r"\D", "", payload["C18"])
        if len(tax) not in (8, 10):
            warnings.append("ІПН / код має незвичну довжину")
            warning_cells.add("C18")

    if payload.get("C29"):
        try:
            year = int(str(payload["C29"]).strip())
            if year < 1900 or year > dt.datetime.now().year + 1:
                warnings.append("Рік випуску поза нормальним діапазоном")
                warning_cells.add("C29")
        except Exception:
            warnings.append("Рік випуску не є числом")
            warning_cells.add("C29")

    if payload.get("C46"):
        try:
            float(str(payload["C46"]).replace(",", "."))
        except Exception:
            errors.append("Ціна з ПДВ — не є числом")
            invalid_cells.add("C46")

    price_no_vat = parse_decimal(payload.get("C44", ""))
    vat = parse_decimal(payload.get("C45", ""))
    price_total = parse_decimal(payload.get("C46", ""))
    if price_no_vat is not None and vat is not None and price_total is not None:
        if abs((price_no_vat + vat) - price_total) > 0.01:
            errors.append("Ціна без ПДВ + ПДВ ≠ Ціна з ПДВ")
            invalid_cells.update({"C44", "C45", "C46"})

    if price_no_vat is not None and price_total is not None and vat is not None and price_total > 0:
        vat_rate = round(vat / price_no_vat * 100, 2) if price_no_vat else None
        if price_no_vat and vat_rate is not None and vat_rate not in (20.0, 7.0, 0.0):
            warnings.append("Ставка ПДВ має незвичне значення")
            warning_cells.add("C45")

    payload["_invalid_cells"] = invalid_cells
    payload["_warning_cells"] = warning_cells
    return payload, errors, warnings


def preview_text_for_act(payload: Dict[str, str]) -> str:
    return "\n".join([
        "ЧОРНОВИК АКТА 6055",
        "",
        f"Номер: {payload.get('Number', '')}",
        f"Дата: {payload.get('Data', '')}",
        f"Покупець: {payload.get('C15', '')}",
        f"Адреса: {payload.get('C16', '')}",
        f"Паспорт: {payload.get('C17', '')}",
        f"ІПН / код: {payload.get('C18', '')}",
        f"Модель: {payload.get('C20', '')}",
        f"Рік: {payload.get('C29', '')}",
        f"Рама / VIN: {payload.get('cuzov', '')}",
        f"Колір: {payload.get('C28', '')}",
        f"Скорочений ПІБ: {payload.get('fio_short', '')}",
    ])


def preview_text_for_contract(payload: Dict[str, str]) -> str:
    lines = ["ЧОРНОВИК ДОГОВОРУ", ""]
    for key in [
        ("Number", "Number"),
        ("Data", "Data"),
        ("FIO", "FIO"),
        ("pasport", "pasport"),
        ("TaxNumber", "TaxNumber"),
        ("BirthDay", "BirthDay"),
        ("adres", "adres"),
        ("decl", "decl"),
        ("model", "model"),
        ("year", "year"),
        ("color", "color"),
        ("numberdv", "numberdv"),
        ("cuzov", "cuzov"),
        ("cub", "cub"),
        ("znak", "znak"),
        ("price", "price"),
        ("sumtext", "sumtext"),
        ("FIO2", "FIO2"),
    ]:
        lines.append(f"{key[0]}: {payload.get(key[1], '')}")
    return "\n".join(lines)


def preview_text_for_vidatkova(payload: Dict[str, str]) -> str:
    return "\n".join([
        "ЧОРНОВИК ВИДАТКОВОЇ",
        "",
        f"Одержувач: {payload.get('C15', '')}",
        f"Дата / номер: {payload.get('A3', '')}",
        f"Марка, модель: {payload.get('C20', '')}",
        f"Номер рами: {payload.get('cuzov', '')}",
        f"Ціна без ПДВ: {payload.get('C44', '')}",
        f"ПДВ: {payload.get('C45', '')}",
        f"Всього: {payload.get('C46', '')}",
        f"Сума прописом: {payload.get('sumtext', '')}",
    ])


_XLS_CELL_RE = re.compile(r"^[A-Z]+\d+$")


def generate_act_xls_from_state(state: Dict[str, str], source_6055: Path, out_path: Path) -> Path:
    shutil.copy2(source_6055, out_path)
    updates = {cell: state.get(cell, "") for cell in FORM_CELLS if _XLS_CELL_RE.match(cell)}
    full_fio = str(state.get("C15", "")).strip()
    short_fio = str(state.get("E15", "")).strip()
    if full_fio and short_fio:
        updates["C15"] = f"{full_fio} / {short_fio}"
    updates["D56"] = short_fio or short_name(full_fio)
    updates.pop("E15", None)  # E15 is app-internal; short name used only in D56
    updates["C43"] = updates.get("C39", "")  # C43 (Номер рами) деривується з VIN
    write_xls_cells(out_path, "Worksheet", updates, backup=False)
    return out_path


def generate_vidatkova_xls_from_state(state: Dict[str, str], template_path: Path, out_path: Path) -> Path:
    payload = parse_state(state)
    shutil.copy2(template_path, out_path)

    discount = parse_decimal(state.get("DISC", "")) or 0.0
    price_no_vat = parse_decimal(payload.get("C44", "")) or 0.0
    vat_amt = parse_decimal(payload.get("C45", "")) or 0.0
    price_total = parse_decimal(payload.get("C46", "")) or 0.0

    if discount > 0:
        price_no_vat_d = max(0.0, price_no_vat - discount)
        vat_d = vat_amt
        price_total_d = max(0.0, price_total - discount)
        sumtext = amount_to_words_uah(price_total_d)
    else:
        price_no_vat_d = price_no_vat
        vat_d = vat_amt
        price_total_d = price_total
        sumtext = payload.get("sumtext", "")

    updates: Dict[str, object] = {
        "C6": payload["FIO"],
        "C7": "той самий",
        "G9": payload["A3"],
        "D10": payload["A3"],
        "B13": payload.get("C21", "МОПЕД"),
        "C13": payload["model"],
        "D13": payload["cuzov"],
        "E13": "шт",
        "F13": 1,
        "G13": price_no_vat if price_no_vat else payload.get("C44", ""),
        "H13": price_no_vat if price_no_vat else payload.get("C44", ""),
        "H15": price_no_vat_d if price_no_vat_d else payload.get("C44", ""),
        "H16": vat_d if vat_d else payload.get("C45", ""),
        "H17": price_total_d if price_total_d else payload.get("C46", ""),
        "A19": sumtext,
        "B20": vat_d if vat_d else payload.get("C45", ""),
        "F23": payload["FIO"],
    }
    if discount > 0:
        updates["H14"] = discount
    write_xls_cells(out_path, "Лист1", updates, backup=False)
    return out_path


def generate_contract_docx_fallback(state: Dict[str, str], out_path: Path) -> "Path | None":
    """Generate full-text contract .docx without Word COM. Returns None if python-docx unavailable."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, Cm  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError:
        return None

    payload = parse_state(state)

    def v(key: str) -> str:
        val = str(payload.get(key, "")).strip()
        return val if val else "___"

    sumtext_plain = contract_sumtext_plain(payload.get("sumtext", "")) or v("sumtext")

    doc = Document()
    # Use Times New Roman 12pt throughout for legal document appearance
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(1.5)

    def add_centered_bold(text: str, size: int = 12) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)

    def add_para(text: str, indent: bool = False) -> None:
        p = doc.add_paragraph(text)
        if indent:
            p.paragraph_format.first_line_indent = Pt(36)

    def add_bold_para(text: str) -> None:
        p = doc.add_paragraph()
        p.add_run(text).bold = True

    # Title
    add_centered_bold("ДОГОВІР КУПІВЛІ-ПРОДАЖУ ТРАНСПОРТНОГО ЗАСОБУ", size=14)
    add_centered_bold(
        f"№ {v('Number')}    місто Вінниця, Україна,    {v('Data')}"
    )
    doc.add_paragraph()

    # Parties
    add_para(
        "Ми, які підписалися нижче:\n"
        "Приватне Підприємство ВКФ МОСТ, код ЄДРПОУ 20112221, м. Вінниця, вул. Івана Богуна, 1/13 "
        "(надалі «Продавець»), в особі директора Колійчука Ігора Миколайовича, який діє на підставі "
        f"Статуту підприємства з однієї сторони та {v('FIO')}, {v('BirthDay')} року народження, "
        f"паспорт серія {v('pasport')}, реєстраційний номер облікової картки платника податків "
        f"{v('TaxNumber')}, який зареєстрований за адресою: {v('adres')}, (надалі «Покупець»), "
        "який діє на підставі власного волевиявлення, з другої сторони уклали цей Договір про таке:"
    )

    # Section 1
    add_bold_para("1. ПРЕДМЕТ ДОГОВОРУ")
    add_para(
        f"1.1. За цим Договором Продавець продає Покупцю, та зобов'язується передати у власність "
        f"Покупця, а Покупець зобов'язується прийняти у власність транспортний засіб, ввезений на "
        f"територію України на підставі митної декларації {v('decl')}, б/використані мотоцикл, "
        f"який має такі характеристики: марка {v('model')}, {v('year')} року випуску, "
        f"колір {v('color')}, номер кузова (шасі, рама) {v('cuzov')}, номер двигуна {v('numberdv')}, "
        f"об'єм двигуна {v('cub')} см\u00b3, транзитний номерний знак виданий торг. орг. {v('znak')}.",
        indent=True,
    )
    add_para(
        "1.2. Транспортний засіб оглянутий Покупцем. Покупець стверджує, що володіє достатньою "
        "інформацією про Транспортний засіб, що набувається, задоволений його технічним станом та "
        "погоджується прийняти його таким, як він є. Під час огляду Транспортного засобу будь-яких "
        "дефектів, недоліків, які перешкоджають використанню Транспортного засобу за цільовим "
        "призначенням, про які не було повідомлено Продавцем не виявлено. Претензій до Продавця "
        "щодо якісних характеристик відчужуваного Транспортного засобу Покупець не має.",
        indent=True,
    )
    add_para(
        "1.3. Продавець підтверджує, що вказаний мотоцикл нікому не проданий, не подарований, "
        "не закладений, не знаходиться під арештом, судова справа по ньому не ведеться та оформлений "
        "в митному відношенні, гарантує достовірність документів на вказаний мотоцикл, законність "
        "права володіння ним, достовірність номерів, вузлів та агрегатів, в подальшому всі претензії "
        "від покупця по вище перерахованих питаннях продавець приймає на себе.",
        indent=True,
    )
    add_para(
        "1.4. Право власності на Транспортний засіб виникає у Покупця з моменту укладання даного "
        "Договору.",
        indent=True,
    )

    # Section 2
    add_bold_para("2. ЦІНА ДОГОВОРУ")
    add_para(
        f"2.1. За домовленістю сторін ціна Транспортного засобу складає {v('price')} "
        f"({sumtext_plain}) гривень 00 копійок.",
        indent=True,
    )
    add_para(
        "2.2. Оплата Транспортного засобу здійснюється шляхом оплати на безготівковий та/або "
        "готівковий рахунок підприємства з відстрочкою платежу 30 днів та підтверджується "
        "видатковою накладною.",
        indent=True,
    )

    # Section 3
    add_bold_para("3. ІНШІ УМОВИ")
    add_para("3.1. Сторони мають права і обов'язки встановлені законодавством.", indent=True)
    add_para("3.2. Договір є укладеним з моменту його підписання Сторонами.", indent=True)
    add_para(
        "3.3. У разі невиконання або неналежного виконання своїх зобов'язань за цим Договором "
        "сторони несуть відповідальність, передбачену чинним законодавством.",
        indent=True,
    )
    add_para(
        "3.4. Договір укладено у двох примірниках, один для Покупця, другий для Продавця.",
        indent=True,
    )

    # Signatures
    doc.add_paragraph()
    add_centered_bold("ПІДПИСИ:")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "ПРОДАВЕЦЬ:"
    tbl.cell(0, 1).text = "ПОКУПЕЦЬ:"
    tbl.cell(1, 0).text = "Директор КОЛІЙЧУК ІГОР МИКОЛАЙОВИЧ"
    tbl.cell(1, 1).text = f"{v('FIO2')}"
    tbl.cell(2, 0).text = "ПП ВКФ МОСТ:  (підпис) ________________________"
    tbl.cell(2, 1).text = "(підпис) ________________________"

    out_docx = out_path.with_suffix(".docx")
    doc.save(str(out_docx))
    return out_docx


def generate_contract_doc_windows_from_state(state: Dict[str, str], dogovir_template: Path, out_path: Path) -> Path:
    payload = parse_state(state)
    if not dogovir_template.exists():
        raise FileNotFoundError(f"Шаблон договору не знайдено: {dogovir_template}")

    word = None
    try:
        import win32com.client as win32  # type: ignore

        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(dogovir_template.resolve()))
        bookmark_map = {
            "Number": "Number",
            "Data": "Data",
            "FIO": "FIO",
            "pasport": "pasport",
            "TaxNumber": "TaxNumber",
            "BirthDay": "BirthDay",
            "adres": "adres",
            "decl": "decl",
            "model": "model",
            "year": "year",
            "color": "color",
            "numberdv": "numberdv",
            "cuzov": "cuzov",
            "cub": "cub",
            "znak": "znak",
            "price": "price",
            "sumtext": "sumtext",
            "FIO2": "FIO2",
        }
        for key, bookmark_name in bookmark_map.items():
            value = payload.get(key, "")
            if key == "sumtext":
                value = contract_sumtext_plain(str(value))
            if doc.Bookmarks.Exists(bookmark_name):
                doc.Bookmarks(bookmark_name).Range.Text = str(value)

        # Also fill via Content Controls (tag or title match)
        try:
            for i in range(1, doc.ContentControls.Count + 1):
                cc = doc.ContentControls(i)
                cc_name = str(cc.Tag or cc.Title or "")
                if cc_name and cc_name in payload:
                    try:
                        cc.Range.Text = str(payload[cc_name])
                    except Exception:
                        pass
        except Exception:
            pass

        # Also fill legacy FormFields (Name match)
        try:
            for i in range(1, doc.FormFields.Count + 1):
                ff = doc.FormFields(i)
                ff_name = str(ff.Name or "")
                if ff_name and ff_name in payload:
                    try:
                        ff.Result = str(payload[ff_name])
                    except Exception:
                        pass
        except Exception:
            pass

        # Also fill via Find & Replace for {{placeholder}} style text
        try:
            find = doc.Content.Find
            find.ClearFormatting()
            for key, value in payload.items():
                val_str = str(value)
                if key == "sumtext":
                    val_str = contract_sumtext_plain(val_str)
                for ph in (f"{{{{{key}}}}}", f"<<{key}>>", f"[{key}]"):
                    find.Execute(
                        FindText=ph, MatchCase=False, MatchWholeWord=False,
                        MatchWildcards=False, MatchSoundsLike=False,
                        MatchAllWordForms=False, Forward=True,
                        Wrap=1, Format=False,
                        ReplaceWith=val_str, Replace=2,
                    )
        except Exception:
            pass

        doc.SaveAs(str(out_path.resolve()))
        doc.Close(SaveChanges=False)
        return out_path
    except Exception as exc:
        raise RuntimeError(
            "Не вдалося сформувати договір через шаблон Word. "
            "Для формату 1:1 потрібен встановлений Microsoft Word і доступний COM."
        ) from exc
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def save_text_preview(path: Path, title: str, body: str) -> Path:
    path.write_text(f"{title}\n\n{body}\n", encoding="utf-8-sig")
    return path


# ---------------------------------------------------------------------------
# App config persistence
# ---------------------------------------------------------------------------
_CONFIG_FILENAME = "jm_config.json"


def _config_path() -> Path:
    return runtime_app_dir() / _CONFIG_FILENAME


def load_app_config() -> dict:
    """Load persisted settings from JSON, returning empty dict on any error."""
    try:
        p = _config_path()
        if p.exists():
            return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {}


def save_app_config(cfg: dict) -> None:
    """Persist settings dict to JSON file next to the app."""
    try:
        _config_path().write_text(
            json.dumps(cfg, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Clipboard text parser
# ---------------------------------------------------------------------------
def _parse_clipboard_to_fields(text: str) -> Dict[str, str]:
    """Extract form field values from free-form clipboard text (customer info)."""
    result: Dict[str, str] = {}
    clean = " ".join(text.split())  # normalise whitespace

    # ПІБ: Ukrainian person name = exactly 3 uppercase Cyrillic words where
    # the LAST word (patronymic) ends with a known Ukrainian/Russian suffix.
    # Prevents matching institution names like "ДЕРЖАВНА МИТНА СЛУЖБА".
    _PATRONYMIC = re.compile(
        r'(?:ОВИЧ|ЕВИЧ|ЄВИЧ|ОВНА|ЄВНА|ЕВНА|ІВНА|ЇВНА|ІЙНА|ЧЕНКО|НЕНКО|ІЄНКО|ЄНКО)$',
        re.IGNORECASE,
    )
    fio_m = re.search(
        r'\b([А-ЯІЇЄҐ\u02BC\'-]{2,}(?:\s+[А-ЯІЇЄҐ\u02BC\'-]{2,}){2})\b',
        clean,
    )
    if fio_m:
        candidate = fio_m.group(1).strip()
        words = candidate.split()
        if (
            len(words) == 3
            and all(re.fullmatch(r'[А-ЯІЇЄҐ\u02BC\'-]{2,}', w, re.IGNORECASE) for w in words)
            and _PATRONYMIC.search(words[-1])
        ):
            result["C15"] = candidate

    # ІПН / реєстраційний номер: exactly 10 digits
    tax_m = re.search(r'(?<!\d)(\d{10})(?!\d)', clean)
    if tax_m:
        result["C18"] = tax_m.group(1)

    # Паспорт: серія (2 кирилічні літери) + 6 цифр [+ дата + видавник]
    ps_m = re.search(
        r'(?:паспорт\s*)?([А-ЯІЇЄҐ]{2})\s*(\d{6})'
        r'(?:\s*,?\s*від\s*(\d{1,2}[./]\d{1,2}[./]\d{2,4})\s*(?:р\.?)?)?'
        r'(?:[,;]?\s*вида(?:ний|ним|но)\s+([^,;\n]+))?',
        clean, re.IGNORECASE,
    )
    if ps_m:
        series, num, date, issuer = ps_m.groups()
        p = f"{series} {num}"
        if date:
            p += f" від {date.replace('/', '.')}"
        if issuer:
            p += f", виданий {issuer.strip().rstrip('.')}"
        result["C17"] = p

    # Дата народження: DD.MM.YYYY р.н. або просто YYYY р.н.
    dob_m = re.search(
        r'(\d{1,2}[./]\d{1,2}[./]\d{4})\s*(?:р\.?н\.?|рік?\s*народж|дата\s*народж)',
        clean, re.IGNORECASE,
    )
    if dob_m:
        result["C12"] = dob_m.group(1).replace("/", ".")
    else:
        yr_m = re.search(
            r'\b(1[89]\d{2}|20[0-2]\d)\s*(?:р\.?н\.?|р\.?\s*народж|рік?\s*народж)',
            clean, re.IGNORECASE,
        )
        if yr_m:
            result["C12"] = yr_m.group(1)

    # Адреса: після ключових слів
    addr_m = re.search(
        r'(?:адреса|адрес[иу]?|зареєстрований\s+за\s+адресою|місце\s+реєстрації)\s*:?\s*([^\n;]+)',
        clean, re.IGNORECASE,
    )
    if addr_m:
        result["C16"] = addr_m.group(1).strip().rstrip(".,;")

    # VIN: стандартний 17-значний код (без літер I, O, Q)
    vin_m = re.search(r'\b([A-HJ-NPR-Z0-9]{17})\b', clean)
    if vin_m:
        result["C39"] = vin_m.group(1)

    # Номер двигуна: після ключового слова
    eng_m = re.search(
        r'(?:двигун[аиу]?|engine)\s*(?:№|#|номер)?\s*:?\s*([A-Za-zА-ЯІЇЄҐа-яіїєґ0-9\-]{4,20})',
        clean, re.IGNORECASE,
    )
    if eng_m:
        result["C41"] = eng_m.group(1).strip()

    return result


def open_file_with_preference(path: Path, editor_path: str = "") -> None:
    if editor_path:
        subprocess.Popen([editor_path, str(path)])
        return
    if IS_WINDOWS:
        os.startfile(str(path))  # type: ignore[attr-defined]
    else:
        subprocess.Popen(["xdg-open", str(path)])


def slugify_case_name(text: str) -> str:
    raw = re.sub(r"\s+", "_", str(text).strip())
    raw = re.sub(r"[^0-9A-Za-zА-Яа-яІіЇїЄєҐґ_\-]", "", raw)
    return raw.strip("_") or "case"


def build_case_folder_name(state: Dict[str, str]) -> str:
    short_fio = short_name(state.get("C15", "")).replace(".", "")
    return slugify_case_name(short_fio)


def ensure_case_dir(base_out_dir: Path, state: Dict[str, str], unique: bool = False) -> Path:
    folder_name = build_case_folder_name(state)
    case_dir = base_out_dir / folder_name
    if unique and case_dir.exists():
        stamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        case_dir = base_out_dir / f"{folder_name}_{stamp}"
    case_dir.mkdir(parents=True, exist_ok=True)
    return case_dir


# ---------------------------------------------------------------------------
# Autocomplete widget
# ---------------------------------------------------------------------------

_BUILTIN_SUGGESTIONS: Dict[str, list] = {
    "C21": ["МОТОЦИКЛ", "МОПЕД", "СКУТЕР", "ТРИЦИКЛ"],
    "C22": ["ЗАГАЛЬНИЙ", "СПОРТИВНИЙ"],
    "C25": ["A1", "A2", "A3", "L1e", "L2e", "L3", "L3e"],
    "C26": ["МОТОЦИКЛ", "МОПЕД", "СКУТЕР"],
    "C27": ["БЕНЗИН", "ДИЗЕЛЬ", "ЕЛЕКТРО"],
}


class SmartEntry(tk.Frame if tk is not None else object):
    """Auto-expanding multiline entry (tk.Text, wrap=word, 1-4 lines) with StringVar sync and autocomplete."""

    def __init__(self, master, cell: str, textvariable: "tk.StringVar", **kwargs):
        kwargs.pop("relief", None)
        kwargs.pop("highlightthickness", None)
        kwargs.pop("bd", None)
        bg = kwargs.pop("bg", "white")
        fg = kwargs.pop("fg", "#374151")
        insert_fg = kwargs.pop("insertbackground", "#111827")
        border = kwargs.pop("highlightbackground", "#d1d5db")
        popup_select_bg = kwargs.pop("popup_select_bg", "#e07b39")
        popup_select_fg = kwargs.pop("popup_select_fg", "white")
        super().__init__(master, bg=bg, highlightbackground=border,
                         highlightthickness=1, **kwargs)
        self._cell = cell
        self._var = textvariable
        self._syncing = False
        self._popup: "tk.Toplevel | None" = None
        self._lb: "tk.Listbox | None" = None
        self._popup_select_bg = popup_select_bg
        self._popup_select_fg = popup_select_fg
        self._fg = fg

        self.columnconfigure(0, weight=1)
        self._text = tk.Text(
            self, height=1, wrap="word", relief="flat", bd=0,
            bg=bg, fg=fg, insertbackground=insert_fg,
            font=("Segoe UI", 10), padx=4, pady=2, undo=True,
        )
        self._text.grid(row=0, column=0, sticky="ew")

        textvariable.trace_add("write", self._on_var_change)
        self._set_text_value(textvariable.get())
        self.after(50, self._adjust_height)

        self._text.bind("<<Modified>>", self._on_text_modified)
        self._text.bind("<KeyRelease>", self._on_key)
        self._text.bind("<FocusOut>", self._close_popup)
        self._text.bind("<Escape>", self._close_popup)
        self._text.bind("<Down>", self._move_down)
        self._text.bind("<Up>", self._move_up)
        self._text.bind("<Tab>", self._on_tab)
        self._text.bind("<Return>", self._on_return)

    def set_bg(self, color: str) -> None:
        self._text.configure(bg=color)
        self.configure(bg=color)
        # Use dark text on light error/warning backgrounds for readability
        if color in ("#fee2e2", "#fef3c7"):
            self._text.configure(fg="#1f2937")
        else:
            self._text.configure(fg=self._fg)

    def _set_text_value(self, value: str) -> None:
        self._text.delete("1.0", "end")
        if value:
            self._text.insert("1.0", str(value))
        self._text.edit_modified(False)
        self._adjust_height()

    def _adjust_height(self) -> None:
        content = self._text.get("1.0", "end-1c")
        if not content:
            self._text.configure(height=1)
            return
        w = self._text.winfo_width()
        chars_per_line = max(20, w // 7) if w > 20 else 50
        wrapped_lines = sum(
            max(1, (len(seg) + chars_per_line - 1) // chars_per_line)
            for seg in content.split("\n")
        )
        self._text.configure(height=min(max(1, wrapped_lines), 4))

    def _on_var_change(self, *_) -> None:
        if self._syncing:
            return
        self._syncing = True
        new_val = self._var.get()
        if self._text.get("1.0", "end-1c") != new_val:
            self._set_text_value(new_val)
        self._syncing = False

    def _on_text_modified(self, *_) -> None:
        if not self._text.edit_modified():
            return
        if self._syncing:
            self._text.edit_modified(False)
            return
        self._syncing = True
        new_val = self._text.get("1.0", "end-1c")
        if self._var.get() != new_val:
            self._var.set(new_val)
        self._text.edit_modified(False)
        self._adjust_height()
        self._syncing = False

    # ── autocomplete ────────────────────────────────────────────────────────
    def _candidates(self) -> list:
        typed = self._var.get().strip().upper()
        all_ = AutocompleteEntry._suggestions.get(self._cell, [])
        if not typed:
            return all_[:4]
        return [v for v in all_ if v.upper().startswith(typed)][:8]

    def _open_popup(self, candidates: list) -> None:
        if not candidates:
            self._close_popup()
            return
        if self._popup is None or not self._popup.winfo_exists():
            self._popup = tk.Toplevel(self._text)
            self._popup.wm_overrideredirect(True)
            self._popup.wm_attributes("-topmost", True)
            self._lb = tk.Listbox(
                self._popup, relief="solid", bd=1,
                selectbackground=self._popup_select_bg,
                selectforeground=self._popup_select_fg,
                font=("Segoe UI", 9), activestyle="none",
            )
            self._lb.pack(fill="both", expand=True)
            self._lb.bind("<ButtonRelease-1>",
                          lambda e: self._pick(self._lb.curselection()))
        self._lb.delete(0, "end")
        for c in candidates:
            self._lb.insert("end", c)
        self.update_idletasks()
        x = self._text.winfo_rootx()
        y = self._text.winfo_rooty() + self._text.winfo_height()
        w = max(self._text.winfo_width(), 200)
        h = min(len(candidates) * 22 + 4, 180)
        self._popup.geometry(f"{w}x{h}+{x}+{y}")
        self._popup.deiconify()

    def _close_popup(self, *_) -> None:
        if self._popup and self._popup.winfo_exists():
            self._popup.withdraw()

    def _on_key(self, event) -> None:
        if event.keysym in ("Down", "Up", "Return", "Escape", "Tab"):
            return
        self._open_popup(self._candidates())

    def _move_down(self, event) -> str:
        if self._popup and self._popup.winfo_viewable() and self._lb:
            cur = self._lb.curselection()
            nxt = (cur[0] + 1) if cur else 0
            if nxt < self._lb.size():
                self._lb.selection_clear(0, "end")
                self._lb.selection_set(nxt)
                self._lb.see(nxt)
        return "break"

    def _move_up(self, event) -> str:
        if self._popup and self._popup.winfo_viewable() and self._lb:
            cur = self._lb.curselection()
            if cur and cur[0] > 0:
                self._lb.selection_clear(0, "end")
                self._lb.selection_set(cur[0] - 1)
                self._lb.see(cur[0] - 1)
        return "break"

    def _on_tab(self, event) -> "str | None":
        if self._popup and self._popup.winfo_viewable() and self._lb:
            cur = self._lb.curselection()
            if not cur and self._lb.size() > 0:
                self._lb.selection_set(0)
                cur = (0,)
            if cur:
                self._pick(cur)
                return "break"
        return None

    def _on_return(self, event) -> str:
        if self._popup and self._popup.winfo_viewable() and self._lb:
            cur = self._lb.curselection()
            if cur:
                self._pick(cur)
        self._close_popup()
        return "break"  # always prevent newline insertion

    def _pick(self, sel) -> None:
        if not sel or self._lb is None:
            return
        value = self._lb.get(sel[0])
        self._syncing = True
        self._var.set(value)
        self._syncing = False
        self._set_text_value(value)
        self._close_popup()
        self._text.focus_set()


class AutocompleteEntry(tk.Entry if tk is not None else object):
    """tk.Entry with a dropdown autocomplete popup powered by a JSON history file."""

    _suggestions: Dict[str, list] = {}
    _db_path: "Path | None" = None

    @classmethod
    def load_db(cls, path: Path) -> None:
        cls._db_path = path
        merged: Dict[str, list] = {k: list(v) for k, v in _BUILTIN_SUGGESTIONS.items()}
        if path.exists():
            try:
                import json as _json
                data = _json.loads(path.read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    for cell, vals in data.items():
                        if isinstance(vals, list):
                            existing = set(merged.get(cell, []))
                            merged.setdefault(cell, [])
                            for val in vals:
                                if val not in existing:
                                    merged[cell].append(val)
                                    existing.add(val)
            except Exception:
                pass
        cls._suggestions = merged

    @classmethod
    def save_db(cls) -> None:
        if cls._db_path is None:
            return
        try:
            import json as _json
            cls._db_path.write_text(
                _json.dumps(cls._suggestions, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        except Exception:
            pass

    @classmethod
    def record(cls, cell: str, value: str) -> None:
        value = str(value).strip()
        if not value:
            return
        lst = cls._suggestions.setdefault(cell, [])
        if value in lst:
            lst.remove(value)
        lst.insert(0, value)
        if len(lst) > 50:
            del lst[50:]

    def __init__(self, master, cell: str, textvariable: tk.StringVar, **kwargs):
        super().__init__(master, textvariable=textvariable, **kwargs)
        self._cell = cell
        self._var = textvariable
        self._popup: "tk.Toplevel | None" = None
        self._lb: "tk.Listbox | None" = None
        self._selecting = False

        self.bind("<KeyRelease>", self._on_key)
        self.bind("<FocusOut>", self._close)
        self.bind("<Escape>", self._close)
        self.bind("<Down>", self._move_down)
        self.bind("<Up>", self._move_up)
        self.bind("<Tab>", self._on_tab)
        self.bind("<Return>", self._on_enter)

    def _candidates(self) -> list:
        typed = self._var.get().strip().upper()
        all_ = self._suggestions.get(self._cell, [])
        if not typed:
            return all_[:4]
        return [v for v in all_ if v.upper().startswith(typed)][:8]

    def _open(self, candidates: list) -> None:
        if not candidates:
            self._close()
            return
        if self._popup is None or not self._popup.winfo_exists():
            self._popup = tk.Toplevel(self)
            self._popup.wm_overrideredirect(True)
            self._popup.wm_attributes("-topmost", True)
            self._lb = tk.Listbox(
                self._popup,
                relief="solid",
                bd=1,
                selectbackground="#e07b39",
                selectforeground="white",
                font=("Segoe UI", 9),
                activestyle="none",
            )
            self._lb.pack(fill="both", expand=True)
            self._lb.bind("<ButtonRelease-1>", lambda e: self._pick(self._lb.curselection()))
        self._lb.delete(0, "end")
        for c in candidates:
            self._lb.insert("end", c)
        self.update_idletasks()
        x = self.winfo_rootx()
        y = self.winfo_rooty() + self.winfo_height()
        w = max(self.winfo_width(), 200)
        h = min(len(candidates) * 22 + 4, 180)
        self._popup.geometry(f"{w}x{h}+{x}+{y}")
        self._popup.deiconify()

    def _close(self, *_) -> None:
        if self._popup and self._popup.winfo_exists():
            self._popup.withdraw()

    def _on_key(self, event) -> None:
        if event.keysym in ("Down", "Up", "Return", "Escape", "Tab"):
            return
        if self._selecting:
            return
        self._open(self._candidates())

    def _move_down(self, event) -> str:
        if self._popup and self._popup.winfo_viewable() and self._lb:
            cur = self._lb.curselection()
            nxt = (cur[0] + 1) if cur else 0
            if nxt < self._lb.size():
                self._lb.selection_clear(0, "end")
                self._lb.selection_set(nxt)
                self._lb.see(nxt)
        return "break"

    def _move_up(self, event) -> str:
        if self._popup and self._popup.winfo_viewable() and self._lb:
            cur = self._lb.curselection()
            if cur and cur[0] > 0:
                self._lb.selection_clear(0, "end")
                self._lb.selection_set(cur[0] - 1)
                self._lb.see(cur[0] - 1)
        return "break"

    def _on_tab(self, event) -> "str | None":
        if self._popup and self._popup.winfo_viewable() and self._lb:
            cur = self._lb.curselection()
            if not cur and self._lb.size() > 0:
                self._lb.selection_set(0)
                cur = (0,)
            if cur:
                self._pick(cur)
                return "break"
        return None

    def _on_enter(self, event) -> "str | None":
        if self._popup and self._popup.winfo_viewable() and self._lb:
            cur = self._lb.curselection()
            if cur:
                self._pick(cur)
                return "break"
        self._close()
        return None

    def _pick(self, sel) -> None:
        if not sel or self._lb is None:
            return
        value = self._lb.get(sel[0])
        self._selecting = True
        self._var.set(value)
        self._selecting = False
        self._close()
        self.icursor("end")
        self.focus_set()


class App:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.app_dir = runtime_app_dir()
        self.resource_dir = runtime_resource_dir()
        self.out_dir = default_output_dir(self.app_dir)

        self.root.title("Japan moto")
        self.root.geometry("1320x920")
        self.root.minsize(1160, 780)
        self.theme_pref = tk.StringVar(value="auto")
        self.theme_mode = detect_theme_mode()
        self.theme = build_theme_palette(self.theme_mode)

        self.source_path = tk.StringVar(value=str(self.resource_dir / "6055.xls"))
        self.moto_path = tk.StringVar(value=str(self.resource_dir / "6055_MOTO_template.xls"))
        self.dogovir_path = tk.StringVar(value=str(self.resource_dir / "DOGOVIR_6055_template.doc"))
        self.vidatkova_path = tk.StringVar(value=str(self.resource_dir / "vidatkova.xls"))
        self.output_dir_var = tk.StringVar(value=str(self.out_dir))
        self.editor_path = tk.StringVar(value="")
        self.open_after_save = tk.BooleanVar(value=True)
        self.use_word_com = tk.BooleanVar(value=True)

        self.state_vars: Dict[str, tk.StringVar] = {}
        self.widgets: Dict[str, tk.Entry] = {}
        self.summary_var = tk.StringVar(value="")
        self.error_var = tk.StringVar(value="")
        self.warning_var = tk.StringVar(value="")
        self.status_var = tk.StringVar(value="Готово")
        self.log_lines: list[str] = []
        self._syncing_form = False
        self._generation_suggested = False

        ac_db_path = self.out_dir / "autocomplete.json"
        AutocompleteEntry.load_db(ac_db_path)

        # Load persisted settings (overrides defaults set above)
        _cfg = load_app_config()
        if _cfg.get("source_path"):
            self.source_path.set(_cfg["source_path"])
        if _cfg.get("moto_path"):
            self.moto_path.set(_cfg["moto_path"])
        if _cfg.get("dogovir_path"):
            self.dogovir_path.set(_cfg["dogovir_path"])
        if _cfg.get("vidatkova_path"):
            self.vidatkova_path.set(_cfg["vidatkova_path"])
        if _cfg.get("output_dir"):
            self.output_dir_var.set(_cfg["output_dir"])
        if _cfg.get("editor_path"):
            self.editor_path.set(_cfg["editor_path"])
        if _cfg.get("theme_pref"):
            self.theme_pref.set(_cfg["theme_pref"])
            _pref = _cfg["theme_pref"]
            if _pref in ("light", "dark"):
                self.theme_mode = _pref
                self.theme = build_theme_palette(_pref)
        self.open_after_save.set(_cfg.get("open_after_save", True))
        self.use_word_com.set(_cfg.get("use_word_com", True))

        self._load_state_from_source()
        self._build_ui()
        self.refresh_validation()

    def _load_state_from_source(self) -> None:
        source = Path(self.source_path.get())
        state = load_form_state(source)
        for _, cell, _, _ in iter_field_specs():
            self.state_vars[cell] = tk.StringVar(value=state.get(cell, ""))
        self.state_vars["E15"].set(short_name(self.state_vars["C15"].get()))

    def _resolve_theme_mode(self) -> str:
        pref = self.theme_pref.get().strip().lower()
        if pref in ("light", "dark"):
            return pref
        return detect_theme_mode()

    def apply_theme_preference(self, *_):
        new_mode = self._resolve_theme_mode()
        self.theme_mode = new_mode
        self.theme = build_theme_palette(new_mode)
        for child in self.root.winfo_children():
            child.destroy()
        self._build_ui()
        self.refresh_validation(silent=True)

    def _build_ui(self) -> None:
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(1, weight=1)
        self.root.rowconfigure(2, weight=0)
        self.root.configure(bg=self.theme["root_bg"])

        header = tk.Frame(self.root, bg=self.theme["header_bg"], padx=10, pady=8)
        header.grid(row=0, column=0, sticky="ew")
        header.columnconfigure(0, weight=1)
        header.columnconfigure(1, weight=0)
        header.columnconfigure(2, weight=0)

        title_frame = tk.Frame(header, bg=self.theme["header_bg"])
        title_frame.grid(row=0, column=0, sticky="w")
        tk.Label(title_frame, text="Japan moto", fg=self.theme["header_fg"], bg=self.theme["header_bg"], font=("Segoe UI", 18, "bold")).pack(anchor="w")
        tk.Label(title_frame, text="Один екран для акта, договору та видаткової", fg=self.theme["header_sub_fg"], bg=self.theme["header_bg"], font=("Segoe UI", 9)).pack(anchor="w")

        toolbar = tk.Frame(header, bg=self.theme["header_bg"])
        toolbar.grid(row=0, column=1, sticky="e")

        self._make_toolbar_button(toolbar, "Акт", lambda: self.open_draft("act"), 0)
        self._make_toolbar_button(toolbar, "Договір", lambda: self.open_draft("contract"), 1)
        self._make_toolbar_button(toolbar, "Видаткова", lambda: self.open_draft("vidatkova"), 2)
        self._make_toolbar_button(toolbar, "↺ Відновити з файлу", self.reload_source, 3)
        self._make_toolbar_button(toolbar, "✖ Очистити", self.clear_form, 4)
        self._make_toolbar_button(toolbar, "⟲ Генерувати", self.generate_all, 5)
        self._make_toolbar_button(toolbar, "↓ Вставити", self.paste_from_clipboard, 6)

        gear_btn = tk.Button(
            header, text="⚙", command=self.open_settings_dialog,
            bg=self.theme["header_bg"], fg=self.theme["header_muted_fg"], activebackground=self.theme["header_active_bg"],
            activeforeground=self.theme["header_fg"], bd=0, relief="flat",
            font=("Segoe UI", 14), padx=8, pady=4, cursor="hand2",
        )
        gear_btn.grid(row=0, column=2, sticky="ne", padx=(0, 4), pady=2)

        main = tk.Frame(self.root, bg=self.theme["surface_bg"])
        main.grid(row=1, column=0, sticky="nsew", padx=12, pady=12)
        main.columnconfigure(0, weight=1)
        main.rowconfigure(0, weight=1)
        self._build_data_tab(main)

        bottom = tk.Frame(self.root, bg=self.theme["status_bg"], padx=12, pady=6)
        bottom.grid(row=2, column=0, sticky="ew")
        bottom.columnconfigure(0, weight=1)
        bottom.columnconfigure(1, weight=1)
        bottom.columnconfigure(2, weight=1)

        tk.Label(bottom, textvariable=self.status_var, bg=self.theme["status_bg"], fg=self.theme["status_fg"], anchor="w").grid(row=0, column=0, sticky="ew")
        tk.Label(bottom, textvariable=self.error_var, bg=self.theme["status_bg"], fg=self.theme["error_fg"], anchor="w", wraplength=640).grid(row=0, column=1, sticky="ew")
        tk.Label(bottom, textvariable=self.warning_var, bg=self.theme["status_bg"], fg=self.theme["warn_fg"], anchor="w", wraplength=640).grid(row=0, column=2, sticky="ew")

        self.write_log(f"Platform: {platform.system()}")
        self.write_log(f"Output folder: {self.out_dir}")
        if not IS_WINDOWS:
            self.write_log("Демо-режим поза Windows: формування договору через Word-шаблон недоступне.")

    def _make_toolbar_button(self, parent, text: str, command, column: int) -> None:
        button = tk.Button(
            parent,
            text=text,
            command=command,
            bg=self.theme["header_bg"],
            fg=self.theme["toolbar_fg"],
            activebackground=self.theme["header_active_bg"],
            activeforeground=self.theme["header_fg"],
            bd=0,
            relief="flat",
            padx=10,
            pady=4,
            cursor="hand2",
        )
        button.grid(row=0, column=column, padx=(0, 6), sticky="e")

    def _add_copy_paste_menu(self, entry) -> None:
        if isinstance(entry, SmartEntry):
            target = entry._text

            def _select_all_fn() -> None:
                target.tag_add("sel", "1.0", "end")
                target.focus_set()

            def _select_all(event: "tk.Event") -> str:
                target.tag_add("sel", "1.0", "end")
                return "break"
        else:
            target = entry

            def _select_all_fn() -> None:
                target.select_range(0, "end")
                target.focus_set()

            def _select_all(event: "tk.Event") -> str:
                event.widget.select_range(0, "end")
                event.widget.icursor("end")
                return "break"

        menu = tk.Menu(target, tearoff=0)
        menu.add_command(label="Вирізати", command=lambda: target.event_generate("<<Cut>>"))
        menu.add_command(label="Копіювати", command=lambda: target.event_generate("<<Copy>>"))
        menu.add_command(label="Вставити", command=lambda: target.event_generate("<<Paste>>"))
        menu.add_separator()
        menu.add_command(label="Виділити все", command=_select_all_fn)

        def _show_menu(event: "tk.Event") -> None:
            menu.tk_popup(event.x_root, event.y_root)

        target.bind("<Button-3>", _show_menu)
        target.bind("<Control-a>", _select_all)

    def _build_data_tab(self, parent: tk.Frame) -> None:
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)

        outer = tk.Frame(parent, bg=self.theme["surface_bg"])
        outer.grid(row=0, column=0, sticky="nsew")
        outer.columnconfigure(0, weight=1)
        outer.rowconfigure(0, weight=1)

        canvas = tk.Canvas(outer, bg=self.theme["surface_bg"], highlightthickness=0)
        _sbar_style = "Themed.Vertical.TScrollbar"
        _st = ttk.Style()
        _st.configure(_sbar_style,
            background=self.theme["scrollbar_bg"],
            troughcolor=self.theme["scrollbar_trough"],
            arrowcolor=self.theme["scrollbar_arrow"],
            bordercolor=self.theme["scrollbar_trough"],
            darkcolor=self.theme["scrollbar_bg"],
            lightcolor=self.theme["scrollbar_bg"],
            borderwidth=0,
            arrowsize=10,
        )
        _st.map(_sbar_style, background=[("active", self.theme["scrollbar_bg"])])
        scrollbar = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview, style=_sbar_style)
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")

        content = tk.Frame(canvas, bg=self.theme["surface_bg"])
        content.columnconfigure(0, weight=1)
        content.columnconfigure(1, weight=1)
        content_win = canvas.create_window((0, 0), window=content, anchor="nw")
        content.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", lambda e: canvas.itemconfigure(content_win, width=e.width))

        def _scroll(e):
            canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")

        def _scroll_up(e):
            canvas.yview_scroll(-1, "units")

        def _scroll_dn(e):
            canvas.yview_scroll(1, "units")

        canvas.bind_all("<MouseWheel>", _scroll)
        canvas.bind_all("<Button-4>", _scroll_up)
        canvas.bind_all("<Button-5>", _scroll_dn)

        last_idx = len(FIELD_SECTIONS) - 1
        for index, (group_name, fields) in enumerate(FIELD_SECTIONS):
            row = index // 2
            col = index % 2
            card = tk.Frame(
                content,
                bg=self.theme["card_bg"],
                highlightbackground=self.theme["card_border"],
                highlightthickness=1,
                padx=14,
                pady=14,
            )
            if index == last_idx and last_idx % 2 == 0:
                card.grid(row=row, column=0, columnspan=2, sticky="nsew", padx=8, pady=8)
            else:
                card.grid(row=row, column=col, sticky="nsew", padx=8, pady=8)
            content.grid_columnconfigure(col, weight=1)
            self._build_field_section(card, group_name, fields)

    def _build_field_section(self, parent, group_name: str, fields):
        parent.columnconfigure(1, weight=1)
        row_offset = 0
        if group_name:
            tk.Label(parent, text=group_name, bg=self.theme["card_bg"], fg=self.theme["section_fg"],
                     font=("Segoe UI", 12, "bold")).grid(
                         row=0, column=0, columnspan=2, sticky="w", pady=(0, 10))
            row_offset = 1
        for row, (cell, label, required) in enumerate(fields):
            actual_row = row + row_offset
            display = f"{label}{' *' if required else ''}"
            tk.Label(parent, text=display, bg=self.theme["card_bg"], fg=self.theme["label_fg"], anchor="w").grid(row=actual_row, column=0, sticky="w", padx=(0, 10), pady=4)

            var = self.state_vars[cell]
            if cell == "E15":
                entry = tk.Entry(parent, textvariable=var, relief="flat", highlightthickness=1, bd=0,
                                 bg=self.theme["readonly_bg"], fg=self.theme["readonly_fg"], insertbackground=self.theme["entry_insert"],
                                 state="readonly", readonlybackground=self.theme["readonly_bg"])
            else:
                entry = SmartEntry(
                    parent,
                    cell=cell,
                    textvariable=var,
                    bg=self.theme["entry_bg"],
                    fg=self.theme["entry_fg"],
                    insertbackground=self.theme["entry_insert"],
                    highlightbackground=self.theme["input_border"],
                    popup_select_bg=self.theme["popup_select_bg"],
                    popup_select_fg=self.theme["popup_select_fg"],
                )
            entry.grid(row=actual_row, column=1, sticky="ew", pady=4)
            self.widgets[cell] = entry

            self._add_copy_paste_menu(entry)
            if cell != "E15":
                var.trace_add("write", self._on_state_change)

    def _build_settings_tab(self, parent: ttk.Frame) -> None:
        parent.columnconfigure(1, weight=1)

        rows = [
            ("6055.xls", self.source_path, True),
            ("6055_MOTO_template.xls", self.moto_path, True),
            ("DOGOVIR_6055_template.doc", self.dogovir_path, True),
            ("vidatkova.xls", self.vidatkova_path, False),
            ("Output folder", tk.StringVar(value=str(self.out_dir)), False),
            ("Editor path", self.editor_path, False),
        ]

        self.output_dir_var = rows[4][1]

        for row, (label, var, editable) in enumerate(rows):
            ttk.Label(parent, text=label).grid(row=row, column=0, sticky="w", padx=(0, 10), pady=5)
            entry = ttk.Entry(parent, textvariable=var)
            entry.grid(row=row, column=1, sticky="ew", pady=5)
            if label in {"6055.xls", "6055_MOTO_template.xls", "DOGOVIR_6055_template.doc", "vidatkova.xls", "Output folder", "Editor path"}:
                ttk.Button(parent, text="Огляд", command=lambda v=var: self.browse_path(v)).grid(row=row, column=2, padx=(8, 0))

        ttk.Checkbutton(parent, text="Відкривати файл після генерації", variable=self.open_after_save).grid(row=6, column=0, columnspan=2, sticky="w", pady=(12, 0))
        ttk.Button(parent, text="Перезавантажити шаблон", command=self.reload_source).grid(row=7, column=0, sticky="w", pady=(12, 0))
        ttk.Button(parent, text="Зберегти новий шаблон", command=self.save_source_changes).grid(row=7, column=1, sticky="w", pady=(12, 0))

    def _build_log_tab(self, parent: ttk.Frame) -> None:
        parent.rowconfigure(0, weight=1)
        parent.columnconfigure(0, weight=1)
        self.log = tk.Text(parent, height=12, wrap="word")
        self.log.grid(row=0, column=0, sticky="nsew")

    def browse_path(self, var: tk.StringVar) -> None:
        if var is self.output_dir_var:
            value = filedialog.askdirectory(initialdir=var.get() or str(self.out_dir)) if filedialog else ""
        else:
            value = filedialog.askopenfilename(initialdir=str(self.app_dir)) if filedialog else ""
        if value:
            var.set(value)

    def open_settings_dialog(self) -> None:
        dialog = tk.Toplevel(self.root)
        dialog.title("Шаблони і вихід")
        dialog.configure(bg=self.theme["card_bg"])
        dialog.transient(self.root)
        dialog.grab_set()

        dialog.columnconfigure(1, weight=1)
        fields = [
            ("Витягнути інфо з файла", self.source_path),
            ("Шаблон акту", self.moto_path),
            ("Шаблон договору", self.dogovir_path),
            ("Шаблон видаткової", self.vidatkova_path),
            ("Папка збереження", self.output_dir_var),
            ("Редактор", self.editor_path),
        ]

        tk.Label(
            dialog,
            text="Параметри шаблонів",
            font=("Segoe UI", 14, "bold"),
            bg=self.theme["card_bg"],
            fg=self.theme["label_fg"],
        ).grid(row=0, column=0, columnspan=3, sticky="w", padx=16, pady=(14, 8))
        for row, (label, var) in enumerate(fields, start=1):
            tk.Label(dialog, text=label, bg=self.theme["card_bg"], fg=self.theme["label_fg"]).grid(row=row, column=0, sticky="w", padx=16, pady=6)
            tk.Entry(dialog, textvariable=var, bg=self.theme["entry_bg"], fg=self.theme["entry_fg"], insertbackground=self.theme["entry_insert"]).grid(row=row, column=1, sticky="ew", pady=6)
            tk.Button(
                dialog,
                text="Огляд",
                command=lambda v=var: self.browse_path(v),
                bg=self.theme["btn_bg"],
                fg=self.theme["btn_fg"],
                activebackground=self.theme["header_active_bg"],
                activeforeground=self.theme["header_fg"],
            ).grid(row=row, column=2, padx=12)

        theme_row = len(fields) + 1
        tk.Label(dialog, text="Тема", bg=self.theme["card_bg"], fg=self.theme["label_fg"]).grid(row=theme_row, column=0, sticky="w", padx=16, pady=6)
        theme_box = tk.Frame(dialog, bg=self.theme["card_bg"])
        theme_box.grid(row=theme_row, column=1, columnspan=2, sticky="w", pady=6)
        for idx, (label, value) in enumerate((("Авто", "auto"), ("Світла", "light"), ("Темна", "dark"))):
            tk.Radiobutton(
                theme_box,
                text=label,
                value=value,
                variable=self.theme_pref,
                bg=self.theme["card_bg"],
                fg=self.theme["label_fg"],
                selectcolor=self.theme["entry_bg"],
                activebackground=self.theme["card_bg"],
                activeforeground=self.theme["label_fg"],
                command=self.apply_theme_preference,
            ).grid(row=0, column=idx, sticky="w", padx=(0, 12))

        tk.Checkbutton(
            dialog,
            text="Відкривати файл після генерації",
            variable=self.open_after_save,
            bg=self.theme["card_bg"],
            fg=self.theme["label_fg"],
            selectcolor=self.theme["entry_bg"],
        ).grid(row=theme_row + 1, column=0, columnspan=3, sticky="w", padx=16, pady=(10, 4))
        tk.Checkbutton(
            dialog,
            text="Інтеграція з Microsoft Office (Word COM)",
            variable=self.use_word_com,
            bg=self.theme["card_bg"],
            fg=self.theme["label_fg"],
            selectcolor=self.theme["entry_bg"],
        ).grid(row=theme_row + 2, column=0, columnspan=3, sticky="w", padx=16, pady=(0, 6))
        tk.Button(
            dialog,
            text="↺ Перезавантажити шаблон",
            command=lambda: [self.reload_source(), dialog.destroy()],
            bg=self.theme["btn_bg"],
            fg=self.theme["btn_fg"],
            activebackground=self.theme["header_active_bg"],
            activeforeground=self.theme["header_fg"],
        ).grid(row=theme_row + 3, column=0, sticky="w", padx=16, pady=8)
        tk.Button(
            dialog,
            text="Зберегти налаштування",
            command=lambda: [self._save_settings(), self.save_source_changes(), dialog.destroy()],
            bg=self.theme["btn_bg"],
            fg=self.theme["btn_fg"],
            activebackground=self.theme["header_active_bg"],
            activeforeground=self.theme["header_fg"],
        ).grid(row=theme_row + 3, column=1, sticky="w", pady=8)
        tk.Button(
            dialog,
            text="Закрити",
            command=lambda: [self._save_settings(), dialog.destroy()],
            bg=self.theme["btn_bg"],
            fg=self.theme["btn_fg"],
            activebackground=self.theme["header_active_bg"],
            activeforeground=self.theme["header_fg"],
        ).grid(row=theme_row + 3, column=2, sticky="e", padx=12, pady=12)
        dialog.protocol("WM_DELETE_WINDOW", lambda: [self._save_settings(), dialog.destroy()])

        # Auto-fit height after all widgets are created
        dialog.update_idletasks()
        _dw = 800
        _dh = max(dialog.winfo_reqheight() + 30, 480)
        _sx = dialog.winfo_screenwidth()
        _sy = dialog.winfo_screenheight()
        _px = max(0, (_sx - _dw) // 2)
        _py = max(0, (_sy - _dh) // 2)
        dialog.geometry(f"{_dw}x{_dh}+{_px}+{_py}")

    def _on_state_change(self, *_):
        if self._syncing_form:
            return
        self._syncing_form = True
        c15_val = self.state_vars["C15"].get()
        self.state_vars["E15"].set(short_name(c15_val))
        _payload, errors, _warnings = self.refresh_validation(silent=True)
        self._syncing_form = False
        if not errors and not self._generation_suggested:
            _countable = [c for c in FORM_CELLS if c not in ("E15", "DISC")]
            _filled = sum(
                1 for c in _countable
                if self.state_vars.get(c, tk.StringVar()).get().strip()
            )
            _pct = _filled / len(_countable) if _countable else 1.0
            if _pct >= 0.9:
                self._generation_suggested = True
                if messagebox and messagebox.askyesno(
                    "Готово до генерації",
                    "Всі обов'язкові поля заповнено!\n\nСгенерувати всі документи зараз?",
                ):
                    self.generate_all()

    def collect_state(self) -> Dict[str, str]:
        state = {cell: var.get().strip() for cell, var in self.state_vars.items()}
        state["E15"] = short_name(state.get("C15", ""))
        return state

    def refresh_validation(self, silent: bool = False):
        state = self.collect_state()
        payload, errors, warnings = validate_state(state)

        invalid_cells: set[str] = payload.pop("_invalid_cells", set())
        warning_cells: set[str] = payload.pop("_warning_cells", set())

        for cell, entry in self.widgets.items():
            if cell == "E15":
                continue
            if cell in invalid_cells:
                bg = "#fee2e2"
            elif cell in warning_cells:
                bg = "#fef3c7"
            else:
                bg = self.theme["entry_bg"]
            if isinstance(entry, SmartEntry):
                entry.set_bg(bg)
            else:
                entry.configure(bg=bg)
                if bg in ("#fee2e2", "#fef3c7"):
                    entry.configure(fg="#1f2937")
                else:
                    entry.configure(fg=self.theme["entry_fg"])

        self.summary_var.set(f"Заповнено: {sum(1 for v in state.values() if v)} полів. Номер: {payload.get('Number', '')} | Дата: {payload.get('Data', '')}")
        self.error_var.set("Помилки: " + ("; ".join(errors) if errors else "немає"))
        self.warning_var.set("Попередження: " + ("; ".join(warnings) if warnings else "немає"))

        if not silent:
            self.write_log(self.error_var.get())
            self.write_log(self.warning_var.get())

        return payload, errors, warnings

    def sync_frame_fields(self) -> None:
        state = self.collect_state()
        frame_value = state.get("C39") or state.get("C42")
        if frame_value:
            self._syncing_form = True
            for cell in ("C39", "C42"):
                self.state_vars[cell].set(frame_value)
            self._syncing_form = False
            self.write_log(f"Синхронізовано VIN/шасі: {frame_value}")
            self.refresh_validation(silent=True)

    def clear_form(self) -> None:
        self._syncing_form = True
        for cell in FORM_CELLS:
            self.state_vars[cell].set("")
        self._syncing_form = False
        self._generation_suggested = False
        self.status_var.set("Форму очищено")
        self.write_log("Форму очищено")
        self.refresh_validation(silent=True)

    def reload_source(self) -> None:
        source = Path(self.source_path.get())
        state = load_form_state(source)
        self._syncing_form = True
        for cell in FORM_CELLS:
            self.state_vars[cell].set(state.get(cell, ""))
        self.state_vars["E15"].set(short_name(self.state_vars["C15"].get()))
        self._syncing_form = False
        self.write_log(f"Перезавантажено дані з {source}")
        self.status_var.set(f"Дані завантажено з {source.name}")
        self.refresh_validation(silent=True)

    def write_log(self, text: str) -> None:
        self.log_lines.append(text)
        if hasattr(self, "log"):
            self.log.insert("end", text + "\n")
            self.log.see("end")

    def save_source_changes(self) -> None:
        source = Path(self.source_path.get())
        state = self.collect_state()
        updates = {cell: state.get(cell, "") for cell in FORM_CELLS if cell != "E15" and _XLS_CELL_RE.match(cell)}
        updates["E15"] = short_name(state.get("C15", ""))
        updates["C39"] = state.get("C39", "")
        updates["C42"] = state.get("C42", "")
        updates["C43"] = state.get("C39", "")  # C43 (Номер рами) auto-derived from VIN
        write_xls_cells(source, "Worksheet", updates, backup=True)
        self.write_log(f"Збережено 6055 у {source}")
        self.status_var.set(f"Збережено у {source.name}")

    def _save_settings(self) -> None:
        """Persist all settings (paths, theme, flags) to JSON config file."""
        save_app_config({
            "source_path": self.source_path.get(),
            "moto_path": self.moto_path.get(),
            "dogovir_path": self.dogovir_path.get(),
            "vidatkova_path": self.vidatkova_path.get(),
            "output_dir": self.output_dir_var.get(),
            "editor_path": self.editor_path.get(),
            "theme_pref": self.theme_pref.get(),
            "open_after_save": self.open_after_save.get(),
            "use_word_com": self.use_word_com.get(),
        })

    def paste_from_clipboard(self) -> None:
        """Parse clipboard text and auto-fill matching form fields."""
        try:
            text = self.root.clipboard_get()
        except Exception:
            self.status_var.set("Буфер порожній або недоступний")
            return
        if not text or not text.strip():
            self.status_var.set("Буфер порожній")
            return
        filled = _parse_clipboard_to_fields(text)
        if not filled:
            self.status_var.set("Не розпізнано жодного поля з буфера обміну")
            return
        self._syncing_form = True
        for cell, value in filled.items():
            if cell in self.state_vars and value:
                self.state_vars[cell].set(value)
        if "C15" in filled:
            self.state_vars["E15"].set(short_name(filled["C15"]))
        self._syncing_form = False
        self.write_log(f"Вставлено з буфера: {', '.join(filled.keys())}")
        self.status_var.set(f"Заповнено {len(filled)} поля(-ів) з буфера")
        self.refresh_validation(silent=True)

    def current_payload(self) -> Dict[str, str]:
        payload, errors, warnings = validate_state(self.collect_state())
        return payload

    def open_draft(self, kind: str) -> None:
        DraftWindow(self, kind)

    def _ensure_output_dir(self) -> Path:
        out_dir = Path(self.output_dir_var.get().strip() or str(self.out_dir))
        out_dir.mkdir(parents=True, exist_ok=True)
        return out_dir

    def save_draft(self, kind: str, open_after: bool = True, output_dir: Path | None = None, use_timestamp: bool = True) -> Path:
        state = self.collect_state()
        payload, errors, warnings = validate_state(state)
        if errors:
            raise ValueError("Потрібно виправити помилки перед збереженням: " + "; ".join(errors))

        out_dir = output_dir or self._ensure_output_dir()
        ts = f"_{dt.datetime.now().strftime('%Y%m%d_%H%M%S')}" if use_timestamp else ""
        doc_num = re.sub(r'[\\/:*?"<>|]', "-", payload.get("Number", "")).strip("-")
        num_part = f" \u2116{doc_num}" if doc_num else ""

        if kind == "act":
            out_path = out_dir / f"Акт{num_part}{ts}.xls"
            generate_act_xls_from_state(state, Path(self.source_path.get()), out_path)
        elif kind == "contract":
            template_doc = Path(self.dogovir_path.get())
            out_doc = out_dir / f"Договір{num_part}{ts}.doc"
            out_path = out_doc
            try:
                if not IS_WINDOWS or not self.use_word_com.get():
                    raise RuntimeError("Word COM disabled or not available")
                out_path = generate_contract_doc_windows_from_state(state, template_doc, out_doc)
            except Exception as exc:
                self.write_log(f"Word COM недоступний: {exc}")

                # Copy template + run-level fill (preserves original formatting)
                non_com = generate_contract_non_com(state, template_doc, out_doc)
                if non_com:
                    out_path = non_com
                else:
                    # Last resort: full text docx generator (different appearance)
                    out_docx = out_doc.with_suffix(".docx")
                    result = generate_contract_docx_fallback(state, out_docx)
                    if result:
                        out_path = result
                    else:
                        out_path = out_doc.with_suffix(".txt")
                        save_text_preview(out_path, "ЧОРНОВИК ДОГОВОРУ", preview_text_for_contract(payload))
        elif kind == "vidatkova":
            out_path = out_dir / f"Видаткова{num_part}{ts}.xls"
            generate_vidatkova_xls_from_state(state, Path(self.vidatkova_path.get()), out_path)
        else:
            raise ValueError(f"Невідомий тип чорновика: {kind}")

        self.write_log(f"Згенеровано: {out_path}")
        self.status_var.set(f"Створено {out_path.name}")
        for cell, value in state.items():
            AutocompleteEntry.record(cell, value)
        AutocompleteEntry.save_db()
        if open_after and self.open_after_save.get():
            try:
                open_file_with_preference(out_path, self.editor_path.get().strip())
            except Exception as exc:
                self.write_log(f"Не вдалося відкрити файл: {exc}")
        return out_path

    def generate_all(self) -> None:
        try:
            state = self.collect_state()
            payload, errors, _ = validate_state(state)
            if errors:
                messagebox.showerror("Помилки валідації", "\n".join(errors))
                return
            base_out_dir = self._ensure_output_dir()
            case_dir = base_out_dir / build_case_folder_name(state)
            if case_dir.exists():
                if not messagebox.askyesno(
                    "Перезаписати?",
                    f"Папка вже існує:\n{case_dir}\n\nПерезаписати всі документи?",
                ):
                    return
            case_dir.mkdir(parents=True, exist_ok=True)
            self.save_draft("act", open_after=False, output_dir=case_dir, use_timestamp=False)
            self.save_draft("contract", open_after=False, output_dir=case_dir, use_timestamp=False)
            self.save_draft("vidatkova", open_after=False, output_dir=case_dir, use_timestamp=False)
            self.status_var.set(f"Усі документи збережено у {case_dir.name}")
            self.write_log(f"Усі документи збережено у {case_dir}")
            if self.open_after_save.get():
                try:
                    open_file_with_preference(case_dir, self.editor_path.get().strip())
                except Exception as exc:
                    self.write_log(f"Не вдалося відкрити папку кейсу: {exc}")
        except Exception as exc:
            self.write_log(f"ERROR: {exc}")
            self.write_log(traceback.format_exc())
            if messagebox:
                messagebox.showerror("Помилка", str(exc))


if tk is not None:
    class DraftWindow(tk.Toplevel):
        _DEST_STEMS = {"act": "6055_akt", "contract": "dogovir", "vidatkova": "vidatkova"}

        def __init__(self, app: App, kind: str):
            super().__init__(app.root)
            self.app = app
            self.kind = kind
            self._temp_path: "Path | None" = None
            self.title(self._title())
            self.transient(app.root)

            self.columnconfigure(0, weight=1)

            header = tk.Frame(self, bg="#0f766e", padx=14, pady=12)
            header.grid(row=0, column=0, sticky="ew")
            header.columnconfigure(0, weight=1)
            tk.Label(header, text=self._title(), fg="white", bg="#0f766e", font=("Segoe UI", 15, "bold")).grid(row=0, column=0, sticky="w")
            tk.Label(header, text="Файл відкрито у системній програмі. Перевірте і збережіть.", fg="#ccfbf1", bg="#0f766e", font=("Segoe UI", 9)).grid(row=1, column=0, sticky="w")

            body = tk.Frame(self, bg="#f4efe7", padx=16, pady=14)
            body.grid(row=1, column=0, sticky="ew")
            body.columnconfigure(0, weight=1)
            self.status_lbl = tk.Label(body, text="Генерація файлу...", bg="#f4efe7", fg="#374151", anchor="w", wraplength=600)
            self.status_lbl.grid(row=0, column=0, sticky="ew")
            self.path_lbl = tk.Label(body, text="", bg="#f4efe7", fg="#6b7280", anchor="w", wraplength=600, font=("Segoe UI", 9))
            self.path_lbl.grid(row=1, column=0, sticky="ew", pady=(4, 0))

            actions = tk.Frame(self, bg="#f4efe7", padx=16, pady=12)
            actions.grid(row=2, column=0, sticky="ew")
            ttk.Button(actions, text="Зберегти в папку клієнта", command=self.save_to_case).pack(side="left", padx=(0, 8))
            ttk.Button(actions, text="Відкрити знову", command=self._open_in_app).pack(side="left", padx=(0, 8))
            ttk.Button(actions, text="Закрити", command=self.destroy).pack(side="left")

            self.update_idletasks()
            _w = 640
            _h = self.winfo_reqheight()
            _px = max(0, app.root.winfo_rootx() + (app.root.winfo_width() - _w) // 2)
            _py = max(0, app.root.winfo_rooty() + (app.root.winfo_height() - _h) // 2)
            self.geometry(f"{_w}x{_h}+{_px}+{_py}")
            self.resizable(False, False)
            self.after(100, self._generate_and_open)

        def _title(self) -> str:
            return {"act": "Акт 6055", "contract": "Договір", "vidatkova": "Видаткова"}[self.kind]

        def _generate_and_open(self) -> None:
            import tempfile
            try:
                tmp_dir = Path(tempfile.mkdtemp(prefix="japan_moto_"))
                self._temp_path = self.app.save_draft(self.kind, open_after=False, output_dir=tmp_dir)
                self.status_lbl.configure(text="Файл відкрито. Перевірте документ у системній програмі, потім збережіть або закрийте.")
                self.path_lbl.configure(text=f"Тимчасовий файл: {self._temp_path}")
                self._open_in_app()
            except Exception as exc:
                self.status_lbl.configure(text=f"Помилка генерації: {exc}", fg="#991b1b")
                self.app.write_log(f"DraftWindow error: {exc}")

        def _open_in_app(self) -> None:
            if self._temp_path and self._temp_path.exists():
                try:
                    open_file_with_preference(self._temp_path, self.app.editor_path.get().strip())
                except Exception as exc:
                    messagebox.showerror("Помилка", str(exc), parent=self)
            else:
                messagebox.showwarning("Немає файлу", "Файл не знайдено. Зачекайте генерацію.", parent=self)

        def save_to_case(self) -> None:
            if not self._temp_path or not self._temp_path.exists():
                messagebox.showwarning("Немає файлу", "Спочатку дочекайтесь генерації.", parent=self)
                return
            state = self.app.collect_state()
            base_out_dir = self.app._ensure_output_dir()
            case_dir = base_out_dir / build_case_folder_name(state)
            payload_s = parse_state(state)
            doc_num_s = re.sub(r'[\\/:*?"<>|]', "-", payload_s.get("Number", "")).strip("-")
            num_part_s = f" №{doc_num_s}" if doc_num_s else ""
            _stem_map = {"act": "Акт", "contract": "Договір", "vidatkova": "Видаткова"}
            dest_name = _stem_map.get(self.kind, self._temp_path.stem) + num_part_s + self._temp_path.suffix
            dest = case_dir / dest_name
            if dest.exists():
                if not messagebox.askyesno("Перезаписати?", f"Файл вже існує:\n{dest}\n\nПерезаписати?", parent=self):
                    return
            case_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy2(self._temp_path, dest)
            messagebox.showinfo("Збережено", f"Файл збережено:\n{dest}", parent=self)
            self.app.write_log(f"Збережено: {dest}")
            self.app.status_var.set(f"Збережено {dest.name}")
else:
    class DraftWindow:
        def __init__(self, *args, **kwargs):
            raise RuntimeError("Tkinter is unavailable in this environment")


def run_demo(app_dir: Path, resource_dir: Path) -> None:
    source = app_dir / "6055.xls"
    moto = resource_dir / "6055_MOTO_template.xls"
    dog = resource_dir / "DOGOVIR_6055_template.doc"

    updates = {
        "A3": "№ 8424/26/009999 від 15 травня 2026 року",
        "C15": "ТЕСТОВИЙ ІВАН ПЕТРОВИЧ",
        "E15": "ТЕСТОВИЙ І. П.",
        "C39": "TESTFRAME123456789",
        "C42": "TESTFRAME123456789",
        "C43": "TESTFRAME123456789",
        "D39": "",
        "D42": "",
        "D43": "",
    }
    write_xls_cells(source, "Worksheet", updates, backup=True)

    out_dir = app_dir / "out"
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = dt.datetime.now().strftime("%Y%m%d_%H%M%S")

    out_act = out_dir / f"6055_MOTO_filled_{ts}.xls"
    transfer_to_act_non_windows(source, moto, out_act)

    if IS_WINDOWS:
        windows_transfer_all(source, moto, dog, out_dir)
    else:
        out_prev = out_dir / f"DOGOVIR_preview_{ts}.txt"
        transfer_to_word_preview(source, out_prev)

    print(f"Demo completed. Output folder: {out_dir}")


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="6055 transfer helper")
    p.add_argument("--demo", action="store_true", help="Run non-interactive demo update + transfer")
    return p.parse_args()


def main() -> int:
    args = parse_args()
    app_dir = runtime_app_dir()
    resource_dir = runtime_resource_dir()

    if args.demo:
        run_demo(app_dir, resource_dir)
        return 0

    if tk is None:
        print("Tkinter is unavailable in this environment.")
        return 1

    root = tk.Tk()
    App(root)
    root.mainloop()
    return 0


if __name__ == "__main__":
    sys.exit(main())
